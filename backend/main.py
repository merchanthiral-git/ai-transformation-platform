"""AI Transformation Platform — FastAPI Backend v3.1
Fixes numpy serialization that causes 500 errors on some Python/FastAPI versions.
"""

import base64
from io import BytesIO

import numpy as np
import pandas as pd
from fastapi import FastAPI, UploadFile, File, HTTPException, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse

from app.store import (
    store, sync_work_design, build_reconstruction, build_ai_priority_scores,
    build_skill_analysis, build_deconstruction_summary, build_workstream_breakdown,
    build_task_portfolio, build_reconstruction_rollup, build_redeployment_detail,
    build_skill_shift_matrix, build_transformation_insights,
    build_transformation_recommendations, build_work_dimensions,
    build_value_model, build_role_evolution, build_capacity_waterfall_data,
    build_auto_change_plan, compute_readiness_score, get_manager_candidates,
    enrich_work_design_defaults, build_operating_model_analysis,
    build_compensation_analysis,
)
from app.helpers import get_series, safe_value_counts, dataframe_to_excel_bytes, empty_bundle
from app.routes_auth import auth_router, project_router

app = FastAPI(title="AI Transformation Platform API", version="3.1")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True, allow_methods=["*"], allow_headers=["*"],
)

app.include_router(auth_router)
app.include_router(project_router)


# ═══════════════════════════════════════════════════════════════════════
# NUMPY-SAFE SERIALIZATION — prevents 500 errors from np.int64 etc.
# ═══════════════════════════════════════════════════════════════════════

def _safe(obj):
    """Recursively convert numpy/pandas types to native Python for JSON."""
    if isinstance(obj, dict):
        return {str(k): _safe(v) for k, v in obj.items()}
    if isinstance(obj, (list, tuple)):
        return [_safe(v) for v in obj]
    if isinstance(obj, (np.integer,)):
        return int(obj)
    if isinstance(obj, (np.floating,)):
        return float(obj)
    if isinstance(obj, (np.bool_,)):
        return bool(obj)
    if isinstance(obj, np.ndarray):
        return obj.tolist()
    if isinstance(obj, pd.Timestamp):
        return str(obj)
    try:
        if pd.isna(obj):
            return None
    except (TypeError, ValueError):
        pass
    return obj


def _j(df, limit=500):
    """Convert DataFrame to list of dicts with all values as strings."""
    if df is None or (isinstance(df, pd.DataFrame) and df.empty):
        return []
    out = df.head(limit).copy()
    for c in out.columns:
        out[c] = out[c].fillna("").astype(str)
    return out.to_dict("records")


def _f(func="All", jf="All", sf="All", cl="All"):
    return {"Function": func, "Job Family": jf, "Sub-Family": sf, "Career Level": cl}


def _job_list_for_model(model_id: str):
    if not model_id or model_id not in store.datasets:
        return []
    data = store.get_filtered_data(model_id, _f())
    jobs = []
    for df in [data.get(k, pd.DataFrame()) for k in ["job_catalog", "work_design", "workforce", "org_design"]]:
        if df is not None and not df.empty and "Job Title" in df.columns:
            jobs.extend([str(x).strip() for x in get_series(df, "Job Title").dropna().unique() if str(x).strip()])
    return sorted(set(jobs))


# ═══════════════════════════════════════════════════════════════════════
# MODELS / UPLOAD / RESET
# ═══════════════════════════════════════════════════════════════════════

@app.get("/api/health")
def health_check():
    """Quick health check — returns OK if backend is running."""
    return {"status": "ok", "datasets": list(store.datasets.keys()), "last_loaded": store.last_loaded_model_id}


@app.get("/api/models")
def list_models():
    try:
        models = store.get_model_ids()
        return _safe({"models": models, "last_loaded": store.last_loaded_model_id})
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(500, detail=str(e))


@app.post("/api/upload")
async def upload_files(files: list[UploadFile] = File(...)):
    fl = [(f.filename, await f.read()) for f in files]
    s = store.process_uploads(fl)
    active_model = store.last_loaded_model_id or ""
    return _safe({
        "sheets_loaded": len(s),
        "summary": _j(s),
        "active_model": active_model,
        "jobs": _job_list_for_model(active_model),
        "models": store.get_model_ids(),
    })


@app.post("/api/reset")
def reset_data():
    store.reset()
    return {"ok": True}


@app.get("/api/template")
def download_template():
    """Generate a multi-tab Excel template with schemas, example data, dropdowns, and Op Model guidance."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    from openpyxl.worksheet.datavalidation import DataValidation

    NAVY, BLUE, TEAL, GREEN, AMBER = "0B1120", "3B82F6", "0891B2", "059669", "D97706"
    WHITE, LTBLUE, LGRAY, MGRAY, DKGRAY = "FFFFFF", "DBEAFE", "F1F5F9", "E2E8F0", "64748B"
    hdr_font = Font(name="Arial", bold=True, color=WHITE, size=10)
    hdr_fill = PatternFill("solid", fgColor=NAVY)
    req_fill = PatternFill("solid", fgColor=LTBLUE)
    opt_fill = PatternFill("solid", fgColor=LGRAY)
    hint_font = Font(name="Arial", italic=True, color=DKGRAY, size=9)
    ex_font = Font(name="Arial", color="8892A8", size=9)
    ex_fill = PatternFill("solid", fgColor="F1F5F9")
    title_font = Font(name="Arial", bold=True, size=14, color=NAVY)
    subtitle_font = Font(name="Arial", size=10, color=DKGRAY)
    guide_title = Font(name="Arial", bold=True, size=12, color=TEAL)
    guide_label = Font(name="Arial", bold=True, size=10, color=NAVY)
    guide_font = Font(name="Arial", size=9, color=DKGRAY)
    thin_border = Border(left=Side(style="thin", color=MGRAY), right=Side(style="thin", color=MGRAY),
                         top=Side(style="thin", color=MGRAY), bottom=Side(style="thin", color=MGRAY))
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    wrap = Alignment(vertical="top", wrap_text=True)

    # ── Tab definitions: columns + example row ──
    TABS = {
        "1. Workforce": {
            "desc": "Employee roster — org hierarchy, comp, classification. Powers: Overview, Org Diagnostics, Compensation, Filters. ★★★",
            "color": BLUE,
            "columns": [
                ("Model ID", True, "Your org/company name"), ("Employee ID", True, "Unique employee ID"), ("Employee Name", True, "Full name"),
                ("Manager ID", False, "Manager's Employee ID"), ("Manager Name", False, "Manager's name"), ("Function ID", False, "Function (Finance, HR...)"),
                ("Job Family", False, "Job family group"), ("Sub-Family", False, "Sub-function"), ("Geography", False, "Country/Region"),
                ("Career Track", False, "Executive/Manager/IC/Analyst"), ("Career Level", False, "Grade code (L2, L5...)"), ("Job Title", False, "Current job title"),
                ("Job Description", False, "Role description"), ("Department", False, "Department name"), ("Org Unit", False, "Business unit"),
                ("FTE", False, "1.0 = full-time"), ("Base Pay", False, "Annual base salary"), ("Total Cash", False, "Total cash comp"),
                ("Hire Date", False, "YYYY-MM-DD"), ("Performance Rating", False, "Rating"), ("Critical Role", False, "Yes/No"),
            ],
            "example": ["Acme_Corp", "E001", "Jane Smith", "", "", "Finance", "Corporate Functions", "FP&A", "US",
                        "Analyst", "L2", "Financial Analyst", "Supports budgeting and reporting", "Finance", "Corporate",
                        1.0, 95000, 102000, "2022-03-15", "Exceeds", "No"],
        },
        "2. Work Design": {
            "desc": "Task-level breakdown per job. Each job's tasks MUST sum to 100% time. Powers: Design Lab, AI Prioritization, Scenarios. ★★★",
            "color": AMBER,
            "columns": [
                ("Model ID", True, "Same Model ID"), ("Function ID", False, "Function"), ("Job Family", False, "Family"), ("Sub-Family", False, "Sub-function"),
                ("Geography", False, "Region"), ("Career Track", False, "Track"), ("Career Level", False, "Level"),
                ("Job Title", True, "Must match Workforce title"), ("Job Description", False, "Role summary"),
                ("Workstream", False, "Workflow group"), ("Task ID", False, "Unique ID (e.g., FA-T1)"), ("Task Name", True, "Descriptive task name"),
                ("Description", False, "What the task involves"), ("AI Impact", False, "High/Moderate/Low"), ("Est Hours/Week", False, "Hours per week"),
                ("Current Time Spent %", False, "% of role time (sum to 100)"), ("Time Saved %", False, "% AI can save"),
                ("Task Type", False, "Repetitive/Variable"), ("Interaction", False, "Independent/Interactive/Collaborative"),
                ("Logic", False, "Deterministic/Probabilistic/Judgment-heavy"), ("Primary Skill", False, "Main skill"), ("Secondary Skill", False, "Supporting skill"),
            ],
            "example": ["Acme_Corp", "Finance", "Corporate Functions", "FP&A", "US", "Analyst", "L2",
                        "Financial Analyst", "Supports budgeting", "Reporting", "FA-T1", "Refresh Management Report Pack",
                        "Update monthly board reports", "High", 4, 10, 5,
                        "Repetitive", "Independent", "Deterministic", "Reporting", "Excel"],
        },
        "3. Operating Model": {
            "desc": "Capability map — layers, scope, hierarchy, ownership. Powers: Structure, Maturity, Decisions, Workflow. ★★★ See guidance below row 7.",
            "color": TEAL,
            "columns": [
                ("Model ID", True, "Same Model ID"), ("Scope", True, "Enterprise or function name (Finance, HR, Technology...)"),
                ("Layer", True, "Governance / Core Components / Shared Services / Enabling / Interface"),
                ("Level 1", False, "Capability area (e.g., Planning)"), ("Level 2", False, "Sub-capability (e.g., Forecasting)"),
                ("Level 3", False, "Process/activity"), ("Level 4", False, "Granular detail"),
                ("Description", False, "Describe — mention 'shared/central' or 'embedded/dedicated' for service model detection"),
                ("Owner", False, "CRITICAL: Capability owner title — enables Decision Rights analysis"),
                ("Function ID", False, "Owning function"), ("Job Family", False, "Related family"), ("Sub-Family", False, "Sub-family"),
                ("Geography", False, "Global / US / UK / India — drives Scope Distribution"), ("Career Track", False, "Track"), ("Career Level", False, "Level"),
            ],
            "example": ["Acme_Corp", "Finance", "Core Components", "Planning", "Forecasting", "Monthly Forecast", "Revenue Build",
                        "Monthly revenue forecasting process — embedded in Finance team", "Director, FP&A",
                        "Finance", "Corporate Functions", "FP&A", "US", "", ""],
        },
        "4. Org Design": {
            "desc": "Org hierarchy for span of control, layer analysis. Can overlap with Workforce. ★★☆",
            "color": GREEN,
            "columns": [
                ("Model ID", True, "Same Model ID"), ("Employee ID", True, "Unique ID"), ("Employee Name", True, "Full name"),
                ("Manager ID", False, "Manager's Employee ID"), ("Manager Name", False, "Manager name"), ("Function ID", False, "Function"),
                ("Job Family", False, "Family"), ("Sub-Family", False, "Sub-function"), ("Geography", False, "Region"),
                ("Career Track", False, "Track"), ("Career Level", False, "Level"), ("Job Title", False, "Title"),
                ("Job Description", False, "Description"), ("Department", False, "Department"), ("Org Unit", False, "Org unit"),
            ],
            "example": ["Acme_Corp", "E001", "Jane Smith", "", "", "Finance", "Corporate Functions", "Leadership",
                        "US", "Executive", "L10", "Chief Financial Officer", "Leads all finance", "Finance", "Enterprise"],
        },
        "5. Job Catalog": {
            "desc": "Standardized job architecture — families, levels, career paths. Powers: Career Architecture, Role Clusters. ★★☆",
            "color": "8B5CF6",
            "columns": [
                ("Model ID", True, "Same Model ID"), ("Job Code", False, "Internal code"), ("Job Title", True, "Standardized title"),
                ("Standard Title", False, "Benchmark title"), ("Function ID", False, "Function"), ("Job Family", False, "Family"),
                ("Sub-Family", False, "Sub-function"), ("Geography", False, "Region"), ("Career Track", False, "Track"),
                ("Career Level", False, "Level"), ("Manager or IC", False, "Manager/IC"), ("Job Description", False, "Description"),
                ("Skills", False, "Semicolon-separated skills"), ("Role Purpose", False, "One-line purpose"),
            ],
            "example": ["Acme_Corp", "FIN-ANL-L2", "Financial Analyst", "Financial Analyst", "Finance",
                        "Corporate Functions", "FP&A", "US", "Analyst", "L2", "IC",
                        "Supports budgeting and reporting", "Excel;Analysis;Reporting;PowerBI", "Provide financial insight and support."],
        },
        "6. Change Management": {
            "desc": "Transformation initiatives, RACI, roadmap milestones. Powers: Roadmap, Risk Analysis, Mobilize. ★★☆",
            "color": "E11D48",
            "columns": [
                ("Model ID", True, "Same Model ID"), ("Function ID", False, "Function"), ("Job Family", False, "Family"),
                ("Sub-Family", False, "Sub-function"), ("Geography", False, "Region"), ("Career Track", False, "Track"), ("Career Level", False, "Level"),
                ("Job Title", False, "Affected title"), ("Task Name", False, "Task being changed"), ("Responsible", False, "RACI: does the work"),
                ("Accountable", False, "RACI: approves"), ("Consulted", False, "RACI: consulted"), ("Informed", False, "RACI: informed"),
                ("Initiative", False, "Initiative name"), ("Owner", False, "Initiative owner"), ("Priority", False, "High/Medium/Low"),
                ("Status", False, "Planned/In Progress/Complete"), ("Wave", False, "Wave 1/Wave 2/etc"), ("Start", False, "YYYY-MM-DD"),
                ("End", False, "YYYY-MM-DD"), ("Milestone", False, "Milestone name"), ("Date", False, "Milestone date"),
                ("Risk", False, "Risk description"), ("Dependency", False, "Dependencies"), ("Notes", False, "Additional notes"),
            ],
            "example": ["Acme_Corp", "Finance", "Corporate Functions", "FP&A", "US", "Analyst", "L2",
                        "Financial Analyst", "Reporting Automation", "Analyst", "Director", "Finance Lead", "CFO",
                        "Reporting Automation", "CFO", "High", "Planned", "Wave 1", "2026-01-01", "2026-03-31",
                        "Kickoff", "2026-01-01", "Resistance to change", "IT resource availability", "Phase 1 pilot"],
        },
        "7. Market Data": {
            "desc": "Compensation benchmarks. Job Match Key must match Workforce Job Titles. Powers: Compensation tab. ★☆☆",
            "color": "7C3AED",
            "columns": [
                ("Model ID", True, "Same Model ID"), ("Source", True, "Survey name"), ("Effective Date", False, "Survey date"),
                ("Currency", False, "USD/EUR/GBP"), ("Job Match Key", True, "Must match Job Title"), ("Survey Job Title", False, "Survey title"),
                ("Function ID", False, "Function"), ("Job Family", False, "Family"), ("Sub-Family", False, "Sub-function"),
                ("Geography", False, "Region"), ("Career Track", False, "Track"), ("Career Level", False, "Level"),
                ("Base 25th", False, "P25 base"), ("Base 50th", False, "P50 base (median)"), ("Base 75th", False, "P75 base"),
                ("TCC 25th", False, "P25 total cash"), ("TCC 50th", False, "P50 total cash"), ("TCC 75th", False, "P75 total cash"),
                ("LTIP 25th", False, "P25 LTIP"), ("LTIP 50th", False, "P50 LTIP"), ("LTIP 75th", False, "P75 LTIP"),
            ],
            "example": ["Acme_Corp", "Radford 2025", "2025-04-01", "USD", "Financial Analyst", "Financial Analyst",
                        "Finance", "Corporate Functions", "FP&A", "US", "Analyst", "L2",
                        80000, 95000, 110000, 85000, 102000, 120000, "", "", ""],
        },
    }

    wb = Workbook()

    # ── Instructions tab ──
    ws = wb.active
    ws.title = "Instructions"
    ws.sheet_properties.tabColor = BLUE
    ws.column_dimensions["A"].width = 90
    ws.column_dimensions["B"].width = 40
    lines = [
        ("AI Transformation Platform — Data Template", title_font),
        ("", subtitle_font),
        ("This workbook has 8 tabs — one for each dataset the platform ingests.", guide_font),
        ("Fill as many as you can. The more data, the richer the analysis.", guide_font),
        ("", guide_font),
        ("HOW TO USE", guide_title),
        ("1. Use a consistent Model ID across ALL tabs (e.g., your company name).", guide_font),
        ("2. Row 2 = column headers. Row 3 = REQUIRED/Optional. Row 4 = field hints.", guide_font),
        ("3. Row 5 = example data (delete before uploading).", guide_font),
        ("4. Fill your data starting from Row 6.", guide_font),
        ("5. Blue header columns are REQUIRED. Dark columns are optional.", guide_font),
        ("6. Job Titles must match across Workforce, Work Design, and Job Catalog.", guide_font),
        ("7. Work Design tasks per job MUST sum to 100% Current Time Spent %.", guide_font),
        ("8. Save and upload via the Upload button in the platform.", guide_font),
        ("", guide_font),
        ("TAB PRIORITY", guide_title),
        ("★★★ Workforce, Work Design, Operating Model — critical for core analysis", guide_font),
        ("★★☆ Org Design, Job Catalog, Change Management — important for depth", guide_font),
        ("★☆☆ Market Data — optional, powers Compensation tab only", guide_font),
        ("", guide_font),
        ("OPERATING MODEL — SPECIAL GUIDANCE", guide_title),
        ("The Operating Model tab is the most nuanced dataset. Key rules:", guide_font),
        ("• Each ROW = one capability, process, or organizational component.", guide_font),
        ("• SCOPE = which domain this belongs to (Enterprise, Finance, HR, Technology, Legal...).", guide_font),
        ("• LAYER = architectural tier: Governance, Core Components, Shared Services, Enabling, Interface.", guide_font),
        ("• LEVELS 1-4 = hierarchical breakdown (L1=Planning > L2=Forecasting > L3=Monthly Forecast).", guide_font),
        ("• OWNER = who owns the capability — CRITICAL for Decision Rights dashboard.", guide_font),
        ("• DESCRIPTION = mention 'shared' or 'embedded' so the platform can detect Service Model.", guide_font),
        ("• GEOGRAPHY = use 'Global' for enterprise-wide items. Drives Scope Distribution chart.", guide_font),
        ("• Aim for 15-30 rows per function and 3-5 Governance rows for meaningful analysis.", guide_font),
        ("• See the Operating Model tab (row 8+) for a detailed field-by-field guide.", guide_font),
    ]
    for ri, (text, font) in enumerate(lines, 1):
        c = ws.cell(row=ri, column=1, value=text)
        c.font = font

    # ── Data tabs ──
    for tab_name, td in TABS.items():
        cols = td["columns"]
        ex = td.get("example", [])
        ws = wb.create_sheet(title=tab_name)
        ws.sheet_properties.tabColor = td["color"]
        ws.merge_cells(f"A1:{get_column_letter(len(cols))}1")
        ws["A1"] = td["desc"]
        ws["A1"].font = subtitle_font
        ws["A1"].alignment = wrap
        ws.row_dimensions[1].height = 30
        # Row 2: headers
        for ci, (cn, req, _hint) in enumerate(cols, 1):
            c = ws.cell(row=2, column=ci, value=cn)
            c.font = hdr_font
            c.fill = PatternFill("solid", fgColor=BLUE) if req else hdr_fill
            c.alignment = center
            c.border = thin_border
        # Row 3: required/optional
        for ci, (cn, req, _) in enumerate(cols, 1):
            c = ws.cell(row=3, column=ci, value="REQUIRED" if req else "Optional")
            c.font = Font(name="Arial", bold=req, size=9, color=BLUE if req else DKGRAY)
            c.fill = req_fill if req else opt_fill
            c.alignment = center
            c.border = thin_border
        # Row 4: hints
        for ci, (_, _, hint) in enumerate(cols, 1):
            c = ws.cell(row=4, column=ci, value=hint)
            c.font = hint_font
            c.alignment = wrap
            c.border = thin_border
        ws.row_dimensions[4].height = 40
        # Row 5: example data
        for ci, val in enumerate(ex, 1):
            c = ws.cell(row=5, column=ci, value=val)
            c.font = ex_font
            c.fill = ex_fill
            c.border = thin_border
        # Empty rows 6-30
        for ri in range(6, 31):
            for ci in range(1, len(cols) + 1):
                c = ws.cell(row=ri, column=ci, value="")
                c.border = thin_border
                if ri % 2 == 0:
                    c.fill = PatternFill("solid", fgColor="F8FAFC")
        # Column widths
        for ci, (cn, _, _) in enumerate(cols, 1):
            ws.column_dimensions[get_column_letter(ci)].width = max(len(cn) + 5, 15)
        ws.freeze_panes = "A6"
        ws.auto_filter.ref = f"A2:{get_column_letter(len(cols))}2"
        # Dropdowns
        for ci, (cn, _, _) in enumerate(cols, 1):
            cl = get_column_letter(ci)
            dv = None
            if cn == "AI Impact": dv = DataValidation(type="list", formula1='"High,Moderate,Low"', allow_blank=True)
            elif cn == "Task Type": dv = DataValidation(type="list", formula1='"Repetitive,Variable"', allow_blank=True)
            elif cn == "Interaction": dv = DataValidation(type="list", formula1='"Independent,Interactive,Collaborative"', allow_blank=True)
            elif cn == "Logic": dv = DataValidation(type="list", formula1='"Deterministic,Probabilistic,Judgment-heavy"', allow_blank=True)
            elif cn == "Priority": dv = DataValidation(type="list", formula1='"High,Medium,Low"', allow_blank=True)
            elif cn == "Status": dv = DataValidation(type="list", formula1='"Planned,In Progress,Complete,On Hold"', allow_blank=True)
            elif cn == "Wave": dv = DataValidation(type="list", formula1='"Wave 1,Wave 2,Wave 3,Wave 4"', allow_blank=True)
            elif cn == "Manager or IC": dv = DataValidation(type="list", formula1='"Manager,IC"', allow_blank=True)
            elif cn == "Critical Role": dv = DataValidation(type="list", formula1='"Yes,No"', allow_blank=True)
            elif cn == "Layer": dv = DataValidation(type="list", formula1='"Governance,Core Components,Shared Services,Enabling,Interface"', allow_blank=True)
            elif cn == "Performance Rating": dv = DataValidation(type="list", formula1='"Exceeds,Meets,Below,1,2,3,4,5"', allow_blank=True)
            if dv:
                dv.error = "Please select from the dropdown"
                ws.add_data_validation(dv)
                dv.add(f"{cl}6:{cl}500")

    # ── Operating Model tab: append guidance section below example ──
    ws_op = wb["3. Operating Model"]
    gr = 8  # start guidance after a gap row
    ws_op.cell(row=gr, column=1, value="OPERATING MODEL — FIELD-BY-FIELD GUIDE").font = guide_title
    ws_op.merge_cells(f"A{gr}:H{gr}")
    guidance = [
        ("Scope", "The domain this capability belongs to. Use 'Enterprise' for firm-wide governance.\nExamples: Enterprise, Finance, HR, Technology, Legal, Investments, Operations, Marketing."),
        ("Layer", "The architectural tier.\n• Governance = oversight bodies, committees, decision forums.\n• Core Components = primary capabilities that deliver the function's mission.\n• Shared Services = capabilities shared across business lines.\n• Enabling = infrastructure, platforms, tools.\n• Interface = touchpoints with the rest of the org (portals, dashboards)."),
        ("Level 1-4", "Hierarchical decomposition of the capability.\nL1 = Capability area (e.g., 'Planning').\nL2 = Sub-capability (e.g., 'Forecasting').\nL3 = Process (e.g., 'Monthly Revenue Forecast').\nL4 = Activity (e.g., 'Collect actuals from GL').\nNot all levels are required — use as many as your detail allows."),
        ("Description", "Free text describing the capability. IMPORTANT: If you mention 'shared', 'central', 'centralized', 'COE', or 'center of excellence', the platform tags it as Shared Service. If you mention 'embedded', 'dedicated', or 'local', it tags as Embedded. This drives the Service Model analysis."),
        ("Owner", "The role title who owns this capability. CRITICAL for Decision Rights.\nExamples: 'CFO', 'Director FP&A', 'CHRO', 'VP Engineering', 'Head of TA'.\nWithout Owner, the Decision Rights tab will show 'Unassigned'."),
        ("Geography", "Use 'Global' for enterprise-wide capabilities, or a specific region (US, UK, India, APAC) for local ones.\nDrives the Scope Distribution analysis (Global vs Local footprint)."),
        ("Best Practices", "• Aim for 15-30 rows per function for meaningful maturity and structure analysis.\n• Include 3-5 Governance rows (committees, forums) to power Decision Rights.\n• Mix Shared Services and Core Components rows for service model split analysis.\n• Fill Owner for every row — it's the #1 field that unlocks the Decisions dashboard.\n• Use consistent Function ID values that match your Workforce and Work Design tabs."),
    ]
    for i, (label, desc) in enumerate(guidance):
        r = gr + 1 + i
        ws_op.cell(row=r, column=1, value=label).font = guide_label
        ws_op.cell(row=r, column=2, value=desc).font = guide_font
        ws_op.cell(row=r, column=2).alignment = wrap
        ws_op.merge_cells(start_row=r, start_column=2, end_row=r, end_column=10)
        ws_op.row_dimensions[r].height = 60

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                             headers={"Content-Disposition": "attachment; filename=AI_Transformation_Template.xlsx"})


# ── Per-module template download ──
MODULE_TAB_MAP = {
    "snapshot": "1. Workforce",
    "jobs": "5. Job Catalog",
    "scan": "2. Work Design",
    "design": "2. Work Design",
    "simulate": "2. Work Design",
    "build": "4. Org Design",
    "plan": "6. Change Management",
    "opmodel": "3. Operating Model",
}

@app.get("/api/template/{module_id}")
def download_module_template(module_id: str):
    """Download a single-tab Excel template for a specific module."""
    tab_name = MODULE_TAB_MAP.get(module_id)
    if not tab_name:
        raise HTTPException(status_code=404, detail=f"No template for module: {module_id}")

    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    from openpyxl.worksheet.datavalidation import DataValidation

    NAVY, BLUE, TEAL, GREEN, AMBER = "0B1120", "3B82F6", "0891B2", "059669", "D97706"
    WHITE, LTBLUE, LGRAY, MGRAY, DKGRAY = "FFFFFF", "DBEAFE", "F1F5F9", "E2E8F0", "64748B"
    hdr_font = Font(name="Arial", bold=True, color=WHITE, size=10)
    hdr_fill = PatternFill("solid", fgColor=NAVY)
    req_fill = PatternFill("solid", fgColor=LTBLUE)
    opt_fill = PatternFill("solid", fgColor=LGRAY)
    hint_font = Font(name="Arial", italic=True, color=DKGRAY, size=9)
    ex_font = Font(name="Arial", color="8892A8", size=9)
    ex_fill = PatternFill("solid", fgColor="F1F5F9")
    subtitle_font = Font(name="Arial", size=10, color=DKGRAY)
    thin_border = Border(left=Side(style="thin", color=MGRAY), right=Side(style="thin", color=MGRAY),
                         top=Side(style="thin", color=MGRAY), bottom=Side(style="thin", color=MGRAY))
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    wrap = Alignment(vertical="top", wrap_text=True)

    # Reuse same tab definitions
    TABS = {
        "1. Workforce": {"desc":"Employee roster — powers Workforce Snapshot, Org Diagnostics, Filters.","color":BLUE,"columns":[("Model ID",True,"Your org/company name"),("Employee ID",True,"Unique employee ID"),("Employee Name",True,"Full name"),("Manager ID",False,"Manager's Employee ID"),("Manager Name",False,"Manager's name"),("Function ID",False,"Function (Finance, HR...)"),("Job Family",False,"Job family group"),("Sub-Family",False,"Sub-function"),("Geography",False,"Country/Region"),("Career Track",False,"Executive/Manager/IC/Analyst"),("Career Level",False,"Grade code"),("Job Title",False,"Current job title"),("Job Description",False,"Role description"),("Department",False,"Department name"),("Org Unit",False,"Business unit"),("FTE",False,"1.0 = full-time"),("Base Pay",False,"Annual base salary"),("Total Cash",False,"Total cash comp"),("Hire Date",False,"YYYY-MM-DD"),("Performance Rating",False,"Rating"),("Critical Role",False,"Yes/No")],"example":["Acme_Corp","E001","Jane Smith","","","Finance","Corporate Functions","FP&A","US","Analyst","L2","Financial Analyst","Supports budgeting","Finance","Corporate",1.0,95000,102000,"2022-03-15","Exceeds","No"]},
        "2. Work Design": {"desc":"Task-level breakdown per job. Each job's tasks MUST sum to 100% time. Powers: Design Lab, AI Scan, Simulator.","color":AMBER,"columns":[("Model ID",True,"Same Model ID"),("Function ID",False,"Function"),("Job Family",False,"Family"),("Sub-Family",False,"Sub-function"),("Geography",False,"Region"),("Career Track",False,"Track"),("Career Level",False,"Level"),("Job Title",True,"Must match Workforce title"),("Job Description",False,"Role summary"),("Workstream",False,"Workflow group"),("Task ID",False,"Unique ID"),("Task Name",True,"Descriptive task name"),("Description",False,"What the task involves"),("AI Impact",False,"High/Moderate/Low"),("Est Hours/Week",False,"Hours per week"),("Current Time Spent %",False,"% of role time (sum to 100)"),("Time Saved %",False,"% AI can save"),("Task Type",False,"Repetitive/Variable"),("Interaction",False,"Independent/Interactive/Collaborative"),("Logic",False,"Deterministic/Probabilistic/Judgment-heavy"),("Primary Skill",False,"Main skill"),("Secondary Skill",False,"Supporting skill")],"example":["Acme_Corp","Finance","Corporate Functions","FP&A","US","Analyst","L2","Financial Analyst","Supports budgeting","Reporting","FA-T1","Refresh Management Report Pack","Update monthly board reports","High",4,10,5,"Repetitive","Independent","Deterministic","Reporting","Excel"]},
        "3. Operating Model": {"desc":"Capability map — layers, scope, hierarchy, ownership. Powers: Operating Model Lab.","color":TEAL,"columns":[("Model ID",True,"Same Model ID"),("Scope",True,"Enterprise or function name"),("Layer",True,"Governance / Core Components / Shared Services / Enabling / Interface"),("Level 1",False,"Capability area"),("Level 2",False,"Sub-capability"),("Level 3",False,"Process/activity"),("Level 4",False,"Granular detail"),("Description",False,"Describe capability"),("Owner",False,"Capability owner title"),("Function ID",False,"Owning function"),("Job Family",False,"Related family"),("Sub-Family",False,"Sub-family"),("Geography",False,"Global/US/UK/India"),("Career Track",False,"Track"),("Career Level",False,"Level")],"example":["Acme_Corp","Finance","Core Components","Planning","Forecasting","Monthly Forecast","Revenue Build","Monthly revenue forecasting process","Director, FP&A","Finance","Corporate Functions","FP&A","US","",""]},
        "4. Org Design": {"desc":"Org hierarchy for span of control, layer analysis.","color":GREEN,"columns":[("Model ID",True,"Same Model ID"),("Employee ID",True,"Unique ID"),("Employee Name",True,"Full name"),("Manager ID",False,"Manager's Employee ID"),("Manager Name",False,"Manager name"),("Function ID",False,"Function"),("Job Family",False,"Family"),("Sub-Family",False,"Sub-function"),("Geography",False,"Region"),("Career Track",False,"Track"),("Career Level",False,"Level"),("Job Title",False,"Title"),("Job Description",False,"Description"),("Department",False,"Department"),("Org Unit",False,"Org unit")],"example":["Acme_Corp","E001","Jane Smith","","","Finance","Corporate Functions","Leadership","US","Executive","L10","Chief Financial Officer","Leads all finance","Finance","Enterprise"]},
        "5. Job Catalog": {"desc":"Standardized job architecture — families, levels, career paths.","color":"8B5CF6","columns":[("Model ID",True,"Same Model ID"),("Job Code",False,"Internal code"),("Job Title",True,"Standardized title"),("Standard Title",False,"Benchmark title"),("Function ID",False,"Function"),("Job Family",False,"Family"),("Sub-Family",False,"Sub-function"),("Geography",False,"Region"),("Career Track",False,"Track"),("Career Level",False,"Level"),("Manager or IC",False,"Manager/IC"),("Job Description",False,"Description"),("Skills",False,"Semicolon-separated skills"),("Role Purpose",False,"One-line purpose")],"example":["Acme_Corp","FIN-ANL-L2","Financial Analyst","Financial Analyst","Finance","Corporate Functions","FP&A","US","Analyst","L2","IC","Supports budgeting","Excel;Analysis;Reporting","Provide financial insight."]},
        "6. Change Management": {"desc":"Transformation initiatives, RACI, roadmap milestones. Powers: Change Planner.","color":"E11D48","columns":[("Model ID",True,"Same Model ID"),("Function ID",False,"Function"),("Job Family",False,"Family"),("Sub-Family",False,"Sub-function"),("Geography",False,"Region"),("Career Track",False,"Track"),("Career Level",False,"Level"),("Job Title",False,"Affected title"),("Task Name",False,"Task being changed"),("Initiative",False,"Initiative name"),("Owner",False,"Initiative owner"),("Priority",False,"High/Medium/Low"),("Status",False,"Planned/In Progress/Complete"),("Wave",False,"Wave 1/Wave 2/etc"),("Start",False,"YYYY-MM-DD"),("End",False,"YYYY-MM-DD"),("Risk",False,"Risk description"),("Notes",False,"Additional notes")],"example":["Acme_Corp","Finance","Corporate Functions","FP&A","US","Analyst","L2","Financial Analyst","Reporting Automation","Reporting Automation","CFO","High","Planned","Wave 1","2026-01-01","2026-03-31","Resistance to change","Phase 1 pilot"]},
    }

    td = TABS.get(tab_name)
    if not td:
        raise HTTPException(status_code=404, detail=f"Tab not found: {tab_name}")

    wb = Workbook()
    ws = wb.active
    ws.title = tab_name
    ws.sheet_properties.tabColor = td["color"]
    cols = td["columns"]
    ex = td.get("example", [])

    ws.merge_cells(f"A1:{get_column_letter(len(cols))}1")
    ws["A1"] = td["desc"]
    ws["A1"].font = subtitle_font
    ws["A1"].alignment = wrap
    ws.row_dimensions[1].height = 30
    for ci, (cn, req, _) in enumerate(cols, 1):
        c = ws.cell(row=2, column=ci, value=cn)
        c.font = hdr_font; c.fill = PatternFill("solid", fgColor=BLUE) if req else hdr_fill; c.alignment = center; c.border = thin_border
    for ci, (_, req, _) in enumerate(cols, 1):
        c = ws.cell(row=3, column=ci, value="REQUIRED" if req else "Optional")
        c.font = Font(name="Arial", bold=req, size=9, color=BLUE if req else DKGRAY); c.fill = req_fill if req else opt_fill; c.alignment = center; c.border = thin_border
    for ci, (_, _, hint) in enumerate(cols, 1):
        c = ws.cell(row=4, column=ci, value=hint); c.font = hint_font; c.alignment = wrap; c.border = thin_border
    ws.row_dimensions[4].height = 40
    for ci, val in enumerate(ex, 1):
        c = ws.cell(row=5, column=ci, value=val); c.font = ex_font; c.fill = ex_fill; c.border = thin_border
    for ri in range(6, 51):
        for ci in range(1, len(cols) + 1):
            c = ws.cell(row=ri, column=ci, value=""); c.border = thin_border
            if ri % 2 == 0: c.fill = PatternFill("solid", fgColor="F8FAFC")
    for ci, (cn, _, _) in enumerate(cols, 1):
        ws.column_dimensions[get_column_letter(ci)].width = max(len(cn) + 5, 15)
    ws.freeze_panes = "A6"
    for ci, (cn, _, _) in enumerate(cols, 1):
        cl = get_column_letter(ci)
        dv = None
        if cn == "AI Impact": dv = DataValidation(type="list", formula1='"High,Moderate,Low"', allow_blank=True)
        elif cn == "Task Type": dv = DataValidation(type="list", formula1='"Repetitive,Variable"', allow_blank=True)
        elif cn == "Interaction": dv = DataValidation(type="list", formula1='"Independent,Interactive,Collaborative"', allow_blank=True)
        elif cn == "Logic": dv = DataValidation(type="list", formula1='"Deterministic,Probabilistic,Judgment-heavy"', allow_blank=True)
        elif cn == "Priority": dv = DataValidation(type="list", formula1='"High,Medium,Low"', allow_blank=True)
        elif cn == "Status": dv = DataValidation(type="list", formula1='"Planned,In Progress,Complete,On Hold"', allow_blank=True)
        elif cn == "Wave": dv = DataValidation(type="list", formula1='"Wave 1,Wave 2,Wave 3,Wave 4"', allow_blank=True)
        elif cn == "Layer": dv = DataValidation(type="list", formula1='"Governance,Core Components,Shared Services,Enabling,Interface"', allow_blank=True)
        if dv:
            dv.error = "Please select from the dropdown"; ws.add_data_validation(dv); dv.add(f"{cl}6:{cl}500")

    safe_name = tab_name.replace(". ", "_").replace(" ", "_")
    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                             headers={"Content-Disposition": f"attachment; filename={safe_name}_Template.xlsx"})


# ═══════════════════════════════════════════════════════════════════════
# FILTERS & JOBS
# ═══════════════════════════════════════════════════════════════════════

@app.get("/api/filter-options")
def get_filter_options(model_id: str, fn: str = Query("All", alias="func"), jf: str = "All", sf: str = "All"):
    from app.helpers import clean_options, build_filter_dimension_source
    if model_id not in store.datasets:
        return {"functions": ["All"], "job_families": ["All"], "sub_families": ["All"], "career_levels": ["All"]}
    store.ensure_model_bundle(model_id)
    b = store.datasets[model_id]
    src = build_filter_dimension_source(b["workforce"], store.build_org_source(model_id),
        store.build_internal_job_catalog(model_id), b["work_design"], b["operating_model"], b["change_management"])
    if src is None or src.empty:
        return {"functions": ["All"], "job_families": ["All"], "sub_families": ["All"], "career_levels": ["All"]}
    def _n(d, pairs):
        for col, val in pairs:
            if val and val != "All" and col in d.columns:
                d = d[get_series(d, col).astype(str) == val]
        return d
    return _safe({
        "functions": clean_options(get_series(src, "Function ID")),
        "job_families": clean_options(get_series(_n(src, [("Function ID", fn)]), "Job Family")),
        "sub_families": clean_options(get_series(_n(src, [("Function ID", fn), ("Job Family", jf)]), "Sub-Family")),
        "career_levels": clean_options(get_series(_n(src, [("Function ID", fn), ("Job Family", jf), ("Sub-Family", sf)]), "Career Level")),
    })


@app.get("/api/job-options")
def get_job_options(model_id: str, fn: str = Query("All", alias="func"), jf: str = "All", sf: str = "All", cl: str = "All"):
    if model_id not in store.datasets:
        return {"jobs": []}
    data = store.get_filtered_data(model_id, _f(fn, jf, sf, cl))
    jobs = []
    for df in [data.get(k, pd.DataFrame()) for k in ["job_catalog", "work_design", "workforce", "org_design"]]:
        if df is not None and not df.empty and "Job Title" in df.columns:
            jobs.extend([str(x).strip() for x in get_series(df, "Job Title").dropna().unique() if str(x).strip()])
    return _safe({"jobs": sorted(set(jobs))})


# ═══════════════════════════════════════════════════════════════════════
# OVERVIEW
# ═══════════════════════════════════════════════════════════════════════

@app.get("/api/overview")
def get_overview(model_id: str, fn: str = Query("All", alias="func"), jf: str = "All", sf: str = "All", cl: str = "All"):
    if model_id not in store.datasets:
        raise HTTPException(404, "Model not found")
    data = store.get_filtered_data(model_id, _f(fn, jf, sf, cl))
    wf, wd, org, jc = data["workforce"], data["work_design"], data["org_design"], data["job_catalog"]
    mgr_df = get_manager_candidates(org)
    avg_span = round(float(pd.to_numeric(get_series(mgr_df, "Direct Reports"), errors="coerce").fillna(0).mean()), 1) if not mgr_df.empty else 0
    high_ai = 0.0
    if not wd.empty:
        wdt = sync_work_design(wd)
        high_count = int(get_series(wdt, "AI Impact").astype(str).eq("High").sum())
        total_count = len(wdt)
        high_ai = round(high_count / total_count * 100, 1) if total_count else 0
    rt, rd = compute_readiness_score(data)
    fc = safe_value_counts(get_series(wf, "Function ID"))
    ai = safe_value_counts(get_series(wd, "AI Impact"))
    cov = {}
    for k, lbl in {"workforce": "Workforce", "job_catalog": "Job Catalog", "work_design": "Work Design", "org_design": "Org Design", "operating_model": "Operating Model", "change_management": "Change Mgmt", "market_data": "Market Data"}.items():
        df = data.get(k, pd.DataFrame())
        cov[k] = {"label": lbl, "ready": not df.empty, "rows": int(len(df))}
    return _safe({
        "kpis": {
            "employees": int(len(wf)) if not wf.empty else int(len(org)),
            "roles": int(len(jc)),
            "tasks_mapped": int(len(wd)),
            "avg_span": float(avg_span),
            "high_ai_pct": float(high_ai),
            "readiness_score": int(rt),
            "readiness_tier": "Advanced" if rt >= 70 else "Emerging" if rt >= 40 else "Early",
        },
        "readiness_dims": {str(k): int(v) for k, v in rd.items()},
        "func_distribution": [{"name": str(k), "value": int(v)} for k, v in fc.items()],
        "ai_distribution": [{"name": str(k), "value": int(v)} for k, v in ai.items()],
        "data_coverage": cov,
    })


# ═══════════════════════════════════════════════════════════════════════
# DIAGNOSE
# ═══════════════════════════════════════════════════════════════════════

@app.get("/api/diagnose/ai-priority")
def get_ai_priority(model_id: str, fn: str = Query("All", alias="func"), jf: str = "All", sf: str = "All", cl: str = "All"):
    wd = store.get_filtered_data(model_id, _f(fn, jf, sf, cl))["work_design"]
    if wd.empty:
        return {"summary": {}, "top10": [], "workstream_impact": []}
    scored = build_ai_priority_scores(wd)
    if scored.empty:
        return {"summary": {}, "top10": [], "workstream_impact": []}
    qw = scored[(scored["Feasibility"] > scored["Feasibility"].median()) & (scored["Time Impact"] > scored["Time Impact"].median())]
    top10 = scored.nlargest(10, "AI Priority")
    ws_impact = scored.groupby("Workstream", as_index=False)["Time Impact"].sum().sort_values("Time Impact", ascending=False)
    return _safe({
        "summary": {
            "tasks_scored": int(len(scored)),
            "quick_wins": int(len(qw)),
            "total_time_impact": round(float(scored["Time Impact"].sum()), 1),
            "avg_risk": round(float(scored["Risk Score"].mean()), 1),
        },
        "top10": _j(top10, 20),
        "workstream_impact": _j(ws_impact, 30),
    })


@app.get("/api/diagnose/skills")
def get_skill_analysis(model_id: str, fn: str = Query("All", alias="func"), jf: str = "All", sf: str = "All", cl: str = "All"):
    wd = store.get_filtered_data(model_id, _f(fn, jf, sf, cl))["work_design"]
    if wd.empty:
        return {"current": [], "future": [], "gap": []}
    recon = build_reconstruction(wd)
    current, future = build_skill_analysis(wd, recon)
    gap = []
    if not current.empty and not future.empty:
        m = current.rename(columns={"Weight": "Current"}).merge(future.rename(columns={"Weight": "Future"}), on="Skill", how="outer").fillna(0)
        m["Delta"] = m["Future"] - m["Current"]
        gap = _j(m.sort_values("Delta"))
    return _safe({"current": _j(current), "future": _j(future), "gap": gap})


@app.get("/api/diagnose/org")
def get_org_diagnostics(model_id: str, fn: str = Query("All", alias="func"), jf: str = "All", sf: str = "All", cl: str = "All"):
    data = store.get_filtered_data(model_id, _f(fn, jf, sf, cl))
    source = data["org_design"] if not data["org_design"].empty else data["workforce"]
    if source.empty:
        return {"empty": True, "kpis": {}, "managers": [], "layers": [], "span_top15": [], "layer_distribution": []}
    mgr_df = get_manager_candidates(source)
    total = int(len(source))
    avg_span = round(float(pd.to_numeric(get_series(mgr_df, "Direct Reports"), errors="coerce").fillna(0).mean()), 1) if not mgr_df.empty else 0
    max_span = int(pd.to_numeric(get_series(mgr_df, "Direct Reports"), errors="coerce").fillna(0).max()) if not mgr_df.empty else 0
    lc = safe_value_counts(get_series(source, "Career Level"))
    mgr_data = []
    span_top15 = []
    if not mgr_df.empty:
        for _, r in mgr_df.nlargest(15, "Direct Reports").iterrows():
            mgr_data.append({"name": str(r.get("Employee Name", "")), "title": str(r.get("Job Title", "")), "reports": int(r.get("Direct Reports", 0))})
            span_top15.append({"Label": str(r.get("Employee Name", "")), "Direct Reports": int(r.get("Direct Reports", 0))})
    return _safe({
        "kpis": {"total": total, "managers": int(len(mgr_df)), "ics": total - int(len(mgr_df)), "avg_span": avg_span, "max_span": max_span, "layers": int(len(lc))},
        "managers": mgr_data, "span_top15": span_top15,
        "layers": [{"name": str(k), "value": int(v)} for k, v in lc.items()],
        "layer_distribution": [{"name": str(k), "value": int(v)} for k, v in lc.items()],
    })


@app.get("/api/diagnose/data-quality")
def get_data_quality(model_id: str):
    if model_id not in store.datasets:
        raise HTTPException(404, "Model not found")
    readiness_df = store.readiness_summary(model_id)
    ready_count = int((readiness_df["Status"] == "Ready").sum()) if not readiness_df.empty else 0
    missing_count = int((readiness_df["Status"] == "Missing").sum()) if not readiness_df.empty else 0
    total_issues = int(readiness_df["Issues"].sum()) if not readiness_df.empty else 0
    avg_completeness = int(readiness_df["Completeness"].mean()) if not readiness_df.empty else 0
    return _safe({
        "summary": {"ready": ready_count, "missing": missing_count, "total_issues": total_issues, "avg_completeness": avg_completeness},
        "readiness": _j(readiness_df), "upload_log": _j(store.upload_log()), "registry": _j(store.model_registry()),
    })


# ═══════════════════════════════════════════════════════════════════════
# DESIGN
# ═══════════════════════════════════════════════════════════════════════

@app.get("/api/design/job-context")
def get_job_context(model_id: str, job: str, fn: str = Query("All", alias="func"), jf: str = "All", sf: str = "All", cl: str = "All", scenario: str = "Balanced"):
    wd = store.get_filtered_data(model_id, _f(fn, jf, sf, cl))["work_design"]
    if wd.empty:
        return {"error": "No work design data", "kpis": {}, "meta": {}, "description": "", "decon_summary": [], "ws_breakdown": [], "ai_distribution": []}
    job_df = sync_work_design(wd[get_series(wd, "Job Title").astype(str) == job].copy())
    if job_df.empty:
        return {"error": f"No tasks for job: {job}", "kpis": {}, "meta": {}, "description": "", "decon_summary": [], "ws_breakdown": [], "ai_distribution": []}
    total_curr = round(float(pd.to_numeric(get_series(job_df, "Est Hours/Week"), errors="coerce").fillna(0).sum()), 1)
    recon = build_reconstruction(job_df, scenario)
    total_fut = round(float(pd.to_numeric(get_series(recon, "Future Hrs"), errors="coerce").fillna(0).sum()), 1) if not recon.empty else 0
    released = round(total_curr - total_fut, 1)
    released_pct = round(released / total_curr * 100, 1) if total_curr else 0
    evolution, evo_detail = build_role_evolution(job_df, recon, build_redeployment_detail(recon) if not recon.empty else pd.DataFrame())
    meta = {}
    for c in ["Function ID", "Job Family", "Sub-Family", "Career Track", "Career Level", "Geography"]:
        if c in job_df.columns:
            v = get_series(job_df, c).dropna().astype(str).str.strip()
            v = v[v != ""]
            if len(v) > 0:
                meta[c] = str(v.iloc[0])
    jd = ""
    if "Job Description" in job_df.columns:
        v = get_series(job_df, "Job Description").dropna().astype(str).str.strip()
        v = v[v != ""]
        jd = str(v.iloc[0]) if len(v) > 0 else ""
    return _safe({
        "kpis": {"hours_week": total_curr, "tasks": int(len(job_df)), "workstreams": int(len(build_workstream_breakdown(job_df))), "released_hrs": released, "released_pct": released_pct, "future_hrs": total_fut, "evolution": str(evolution)},
        "meta": meta, "description": jd,
        "decon_summary": _j(build_deconstruction_summary(job_df)),
        "ws_breakdown": _j(build_workstream_breakdown(job_df)),
        "ai_distribution": [{"name": str(k), "value": int(v)} for k, v in safe_value_counts(get_series(job_df, "AI Impact")).items()],
        "evolution_detail": evo_detail if isinstance(evo_detail, list) else [],
    })


@app.get("/api/design/deconstruction")
def get_deconstruction(model_id: str, job: str, fn: str = Query("All", alias="func"), jf: str = "All", sf: str = "All", cl: str = "All"):
    wd = store.get_filtered_data(model_id, _f(fn, jf, sf, cl))["work_design"]
    if wd.empty:
        return {"tasks": [], "dimensions": [], "ai_priority": []}
    job_df = sync_work_design(wd[get_series(wd, "Job Title").astype(str) == job].copy())
    if job_df.empty:
        return {"tasks": [], "dimensions": [], "ai_priority": []}
    cols = ["Task ID", "Task Name", "Workstream", "AI Impact", "Est Hours/Week", "Current Time Spent %", "Time Saved %", "Task Type", "Interaction", "Logic", "Primary Skill", "Secondary Skill"]
    edit_df = job_df[[c for c in cols if c in job_df.columns]].copy()
    for c in cols:
        if c not in edit_df.columns:
            edit_df[c] = ""
    if "Task ID" not in edit_df.columns or edit_df["Task ID"].astype(str).str.strip().eq("").all():
        edit_df["Task ID"] = [f"T{i+1}" for i in range(len(edit_df))]
    return _safe({"tasks": _j(edit_df, 200), "dimensions": _j(build_work_dimensions(job_df), 200), "ai_priority": _j(build_ai_priority_scores(job_df), 50)})


@app.post("/api/design/reconstruct")
def compute_reconstruction_endpoint(payload: dict):
    tasks = payload.get("tasks", [])
    scenario = payload.get("scenario", "Balanced")
    if not tasks:
        return {"reconstruction": [], "rollup": [], "value_model": {}, "recommendations": []}
    job_df = sync_work_design(pd.DataFrame(tasks))
    for c in ["Est Hours/Week", "Current Time Spent %", "Time Saved %"]:
        if c in job_df.columns:
            job_df[c] = pd.to_numeric(job_df[c], errors="coerce").fillna(0)
    recon = build_reconstruction(job_df, scenario)
    if recon.empty:
        return {"reconstruction": [], "rollup": [], "value_model": {}, "recommendations": []}
    rd = build_redeployment_detail(recon)
    ac = safe_value_counts(get_series(recon, "Action")) if "Action" in recon.columns else pd.Series()
    wf_c, wf_a, wf_au, wf_o, wf_r = build_capacity_waterfall_data(recon)
    evolution, evo_detail = build_role_evolution(job_df, recon, rd)
    return _safe({
        "reconstruction": _j(recon, 200), "rollup": _j(build_reconstruction_rollup(recon)),
        "redeployment": _j(rd, 200), "insights": _j(build_transformation_insights(job_df, recon, rd)),
        "recommendations": build_transformation_recommendations(job_df, recon, rd),
        "value_model": build_value_model(job_df, recon),
        "evolution": str(evolution), "evolution_detail": evo_detail if isinstance(evo_detail, list) else [],
        "action_mix": {str(k): int(v) for k, v in ac.items()},
        "waterfall": {"current": float(wf_c), "automated": float(wf_a), "augmented": float(wf_au), "other": float(wf_o), "redeployed": float(wf_r)},
    })


# ═══════════════════════════════════════════════════════════════════════
# SIMULATE
# ═══════════════════════════════════════════════════════════════════════

@app.get("/api/simulate/scenarios")
def get_scenarios(model_id: str, fn: str = Query("All", alias="func"), jf: str = "All", sf: str = "All", cl: str = "All"):
    wd = store.get_filtered_data(model_id, _f(fn, jf, sf, cl))["work_design"]
    if wd.empty:
        return {"scenarios": {}}
    results = {}
    for sc in ["Conservative", "Balanced", "Transformative"]:
        t = sync_work_design(enrich_work_design_defaults(wd.copy(), sc))
        ch = float(pd.to_numeric(get_series(t, "Est Hours/Week"), errors="coerce").fillna(0).sum())
        sp = float(pd.to_numeric(get_series(t, "Time Saved %"), errors="coerce").fillna(0).sum())
        tp = float(pd.to_numeric(get_series(t, "Current Time Spent %"), errors="coerce").fillna(0).sum())
        rel = sp / tp * ch if tp else 0
        results[sc] = {
            "current_hrs": round(ch, 1), "future_hrs": round(ch - rel, 1), "released_hrs": round(rel, 1),
            "time_saved_pct": round(sp / tp * 100, 1) if tp else 0,
            "high_ai_tasks": int((get_series(t, "AI Impact").astype(str) == "High").sum()),
        }
    return _safe({"scenarios": results})


@app.get("/api/simulate/readiness")
def get_readiness(model_id: str, fn: str = Query("All", alias="func"), jf: str = "All", sf: str = "All", cl: str = "All"):
    data = store.get_filtered_data(model_id, _f(fn, jf, sf, cl))
    total, dims = compute_readiness_score(data)
    return _safe({
        "score": int(total), "total": int(total),
        "tier": "Advanced" if total >= 70 else "Emerging" if total >= 40 else "Early",
        "dimensions": {str(k): int(v) for k, v in dims.items()},
        "dims": {str(k): int(v) for k, v in dims.items()},
    })


# ═══════════════════════════════════════════════════════════════════════
# MOBILIZE
# ═══════════════════════════════════════════════════════════════════════

@app.get("/api/mobilize/roadmap")
def get_roadmap(model_id: str, fn: str = Query("All", alias="func"), jf: str = "All", sf: str = "All", cl: str = "All"):
    data = store.get_filtered_data(model_id, _f(fn, jf, sf, cl))
    ch, wd, org, op = data["change_management"], data["work_design"], data["org_design"], data["operating_model"]
    auto = build_auto_change_plan(wd, org, op)
    rm = auto if ch.empty else ch
    if rm.empty:
        return {"roadmap": [], "summary": {}, "priority_distribution": {}, "wave_distribution": {}}
    pd_counts = safe_value_counts(get_series(rm, "Priority")) if "Priority" in rm.columns else pd.Series()
    wd_counts = safe_value_counts(get_series(rm, "Wave")) if "Wave" in rm.columns else pd.Series()
    high_pri = int(get_series(rm, "Priority").astype(str).isin(["High", "Critical"]).sum()) if "Priority" in rm.columns else 0
    wave_set = set(get_series(rm, "Wave").dropna().astype(str).str.strip().unique()) if "Wave" in rm.columns else set()
    return _safe({
        "roadmap": _j(rm),
        "summary": {"total": int(len(rm)), "high_priority": high_pri, "waves": int(len(wave_set)), "source": "uploaded" if not ch.empty else "auto"},
        "priority_distribution": {str(k): int(v) for k, v in pd_counts.items()},
        "wave_distribution": {str(k): int(v) for k, v in wd_counts.items()},
    })


@app.get("/api/mobilize/risk")
def get_risk(model_id: str, fn: str = Query("All", alias="func"), jf: str = "All", sf: str = "All", cl: str = "All"):
    wd = store.get_filtered_data(model_id, _f(fn, jf, sf, cl))["work_design"]
    if wd.empty:
        return {"empty": True, "high_risk_tasks": [], "risk_by_workstream": [], "summary": {}}
    scored = build_ai_priority_scores(wd)
    if scored.empty:
        return {"empty": True, "high_risk_tasks": [], "risk_by_workstream": [], "summary": {}}
    hr = scored[scored["Risk Score"] > 4]
    na = scored[(scored["Logic"] == "Judgment-heavy") | (scored["Interaction"] == "Collaborative")]
    rws = scored.groupby("Workstream", as_index=False)["Risk Score"].mean().sort_values("Risk Score", ascending=False)
    return _safe({
        "high_risk_tasks": _j(hr, 30), "risk_by_workstream": _j(rws),
        "summary": {"high_risk_count": int(len(hr)), "no_automate_count": int(len(na)), "avg_risk": round(float(scored["Risk Score"].mean()), 1), "total_assessed": int(len(scored))},
    })


# ═══════════════════════════════════════════════════════════════════════
# OPERATING MODEL
# ═══════════════════════════════════════════════════════════════════════

@app.get("/api/operating-model")
def get_operating_model(model_id: str, fn: str = Query("All", alias="func"), jf: str = "All", sf: str = "All", cl: str = "All"):
    if model_id not in store.datasets:
        return _safe({"kpis": {}, "maturity": [], "structure": [], "workflow": [], "decisions": [], "insights": [], "layer_agg": [], "service_split": [], "scope_dist": [], "decision_load": [], "stage_throughput": []})
    data = store.get_filtered_data(model_id, _f(fn, jf, sf, cl))
    try:
        return _safe(build_operating_model_analysis(data))
    except Exception as e:
        print(f"[operating-model] Error: {e}")
        return _safe({"kpis": {}, "maturity": [], "structure": [], "workflow": [], "decisions": [], "insights": [], "layer_agg": [], "service_split": [], "scope_dist": [], "decision_load": [], "stage_throughput": []})


# ═══════════════════════════════════════════════════════════════════════
# COMPENSATION
# ═══════════════════════════════════════════════════════════════════════

@app.get("/api/compensation")
def get_compensation(model_id: str, fn: str = Query("All", alias="func"), jf: str = "All", sf: str = "All", cl: str = "All"):
    if model_id not in store.datasets:
        return _safe({"kpis": {}, "positioning": [], "by_function": [], "by_level": [], "pay_ranges": [], "insights": [], "detail": []})
    data = store.get_filtered_data(model_id, _f(fn, jf, sf, cl))
    try:
        return _safe(build_compensation_analysis(data))
    except Exception as e:
        print(f"[compensation] Error: {e}")
        return _safe({"kpis": {}, "positioning": [], "by_function": [], "by_level": [], "pay_ranges": [], "insights": [], "detail": []})


# ═══════════════════════════════════════════════════════════════════════
# EXPORT
# ═══════════════════════════════════════════════════════════════════════

@app.get("/api/export/datasets")
def get_export_datasets(model_id: str, fn: str = Query("All", alias="func"), jf: str = "All", sf: str = "All", cl: str = "All"):
    data = store.get_filtered_data(model_id, _f(fn, jf, sf, cl))
    labels = {"workforce": "Workforce", "job_catalog": "Job Catalog", "work_design": "Work Design", "org_design": "Org Design", "operating_model": "Operating Model", "change_management": "Change Mgmt", "market_data": "Market Data"}
    exports = {}
    total_rows = 0
    available_count = 0
    for name, df in data.items():
        avail = not df.empty
        rows = int(len(df))
        if avail:
            available_count += 1
            total_rows += rows
        exports[name] = {"label": labels.get(name, name), "available": avail, "rows": rows, "cols": int(len(df.columns)) if avail else 0, "preview": _j(df, 5)}
    return _safe({"exports": exports, "summary": {"available": available_count, "total_rows": total_rows, "model_id": model_id}})


@app.get("/api/export/download/{dataset_name}")
def download_dataset(dataset_name: str, model_id: str, fn: str = Query("All", alias="func"), jf: str = "All", sf: str = "All", cl: str = "All"):
    data = store.get_filtered_data(model_id, _f(fn, jf, sf, cl))
    if dataset_name not in data or data[dataset_name].empty:
        raise HTTPException(404, "Dataset empty")
    return StreamingResponse(BytesIO(dataframe_to_excel_bytes(data[dataset_name], dataset_name[:31])),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={model_id}_{dataset_name}.xlsx"})


@app.get("/")
def root():
    return {"ok": True, "app": "ai-transformation-platform", "version": "3.1"}



# ═══════════════════════════════════════════════════════════════════════
# AI ASSISTANT — Routes AI requests through Gemini
# ═══════════════════════════════════════════════════════════════════════

import httpx

GEMINI_API_KEY = "AIzaSyDpnXVJz5C_EouGxM0edhSGvYyAprpB5TI"
GEMINI_URL = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={GEMINI_API_KEY}"

@app.get("/api/ai/health")
async def ai_health():
    """Quick test of Gemini connectivity."""
    try:
        async with httpx.AsyncClient(timeout=10.0) as client:
            resp = await client.post(GEMINI_URL, json={
                "contents": [{"parts": [{"text": "Reply with only the word OK"}]}],
                "generationConfig": {"maxOutputTokens": 10}
            })
            data = resp.json()
            if resp.status_code != 200:
                return {"status": "error", "code": resp.status_code, "detail": data.get("error", {}).get("message", "Unknown error")}
            text = data.get("candidates", [{}])[0].get("content", {}).get("parts", [{}])[0].get("text", "")
            return {"status": "ok", "response": text.strip()}
    except Exception as e:
        return {"status": "error", "detail": str(e)}

@app.post("/api/ai/generate")
async def ai_generate(payload: dict):
    """Proxy AI requests to Google Gemini."""
    system_prompt = payload.get("system", "")
    user_message = payload.get("message", "")
    
    gemini_body = {
        "contents": [{"parts": [{"text": f"{system_prompt}\n\n{user_message}"}]}],
        "generationConfig": {
            "temperature": 0.7,
            "maxOutputTokens": 2000,
        }
    }
    
    try:
        async with httpx.AsyncClient(timeout=60.0) as client:
            resp = await client.post(GEMINI_URL, json=gemini_body)
            data = resp.json()
            
            # Extract text from Gemini response
            text = ""
            candidates = data.get("candidates", [])
            if candidates:
                parts = candidates[0].get("content", {}).get("parts", [])
                if parts:
                    text = parts[0].get("text", "")
            
            return {"text": text}
    except Exception as e:
        return {"text": f"AI service error: {str(e)}", "error": True}


# ═══════════════════════════════════════════════════
# SKILLS & TALENT ENDPOINTS
# ═══════════════════════════════════════════════════

# Simple response cache for expensive chained computations
_api_cache = {}
def _cache_key(endpoint, model_id):
    return f"{endpoint}:{model_id}"
def _get_cached(endpoint, model_id):
    key = _cache_key(endpoint, model_id)
    if key in _api_cache:
        return _api_cache[key]
    return None
def _set_cached(endpoint, model_id, data):
    _api_cache[_cache_key(endpoint, model_id)] = data
    # Limit cache size
    if len(_api_cache) > 50:
        oldest = list(_api_cache.keys())[0]
        del _api_cache[oldest]

@app.get("/api/cache/clear")
async def clear_cache():
    _api_cache.clear()
    return {"status": "cleared"}

@app.get("/api/skills/inventory/{model_id}")
async def get_skills_inventory(model_id: str, func: str = "All", jf: str = "All"):
    """Skills Inventory — reads from uploaded Skills tab or generates from work design tasks."""
    ds = store.datasets.get(model_id)
    if not ds:
        return _safe({"employees": [], "skills": [], "clusters": {}, "records": [], "coverage": 0, "sample": True, "total_headcount": 0, "validation_errors": []})
    
    validation_errors = []
    
    # 1. Try real Skills tab first
    skills_df = None
    for key in ds:
        if "skill" in key.lower() and "gap" not in key.lower():
            skills_df = ds[key]
            break
    
    # 2. Get workforce/headcount data
    wf_df = None
    for key in ["workforce", "Workforce", "headcount", "Headcount", "employees", "Employees"]:
        if key in ds and not ds[key].empty:
            wf_df = ds[key]
            break
    
    # 3. Get work design data for skill extraction
    wd_df = None
    for key in ds:
        kl = key.lower()
        if any(t in kl for t in ["work_design", "work design", "tasks", "task", "deconstruction"]):
            wd_df = ds[key]
            break
    
    # Determine employees list
    employees = []
    if wf_df is not None and not wf_df.empty:
        emp_col = next((c for c in wf_df.columns if any(x in c.lower() for x in ["employee", "name", "worker"])), wf_df.columns[0])
        employees = [str(e) for e in wf_df[emp_col].dropna().unique().tolist()]
        role_col = next((c for c in wf_df.columns if any(x in c.lower() for x in ["role", "job", "title", "position"])), None)
    
    # CASE A: Real skills data exists
    if skills_df is not None and not skills_df.empty:
        # Validate columns
        cols = [c.lower() for c in skills_df.columns]
        emp_col = next((c for c in skills_df.columns if any(x in c.lower() for x in ["employee", "name", "worker", "id"])), None)
        skill_col = next((c for c in skills_df.columns if "skill" in c.lower() and "proficiency" not in c.lower() and "level" not in c.lower()), None)
        prof_col = next((c for c in skills_df.columns if any(x in c.lower() for x in ["proficiency", "level", "score", "rating"])), None)
        
        if not emp_col:
            validation_errors.append("Skills tab is missing an Employee/Name column — add it to enable inventory")
        if not skill_col:
            validation_errors.append("Skills tab is missing a Skill Name column — add it to identify skills")
        if not prof_col:
            validation_errors.append("Skills tab is missing a Proficiency/Level column — add it to enable gap analysis")
        
        if emp_col and skill_col:
            records = []
            for _, row in skills_df.iterrows():
                emp = str(row.get(emp_col, "")).strip()
                skill = str(row.get(skill_col, "")).strip()
                prof = 2  # default
                if prof_col:
                    try:
                        prof = int(float(row.get(prof_col, 2)))
                    except:
                        prof = 2
                prof = max(1, min(4, prof))
                if emp and skill:
                    records.append({"employee": emp, "skill": skill, "proficiency": prof})
            
            all_employees = list(set(r["employee"] for r in records))
            all_skills = list(set(r["skill"] for r in records))
            
            # Auto-cluster skills
            clusters = _auto_cluster_skills(all_skills)
            
            # Dedup suggestions
            dedup = _find_skill_duplicates(all_skills)
            
            coverage = round(len(all_employees) / max(len(employees) if employees else len(all_employees), 1) * 100)
            
            return _safe({
                "employees": all_employees[:100],
                "skills": all_skills,
                "clusters": clusters,
                "records": records,
                "coverage": min(coverage, 100),
                "sample": False,
                "total_headcount": len(employees) if employees else len(all_employees),
                "validation_errors": validation_errors,
                "dedup_suggestions": dedup,
            })
    
    # CASE B: No skills tab — extract skills from work design tasks
    if wd_df is not None and not wd_df.empty:
        from app.store import sync_work_design, build_skill_analysis, build_reconstruction
        t = sync_work_design(wd_df)
        current_skills, _ = build_skill_analysis(wd_df)
        
        if not current_skills.empty:
            all_skills = current_skills["Skill"].tolist()
            
            # Generate proficiency records from work design weights
            import random
            random.seed(hash(model_id) % 2**31)
            
            emps = employees[:50] if employees else [f"Employee_{i+1}" for i in range(20)]
            records = []
            for emp in emps:
                for _, srow in current_skills.iterrows():
                    skill = str(srow["Skill"])
                    weight = float(srow.get("Weight", 10))
                    # Higher weight skills = higher baseline proficiency
                    base = 2 if weight > 15 else 1
                    prof = min(4, max(1, base + random.choices([-1, 0, 0, 1], weights=[10, 40, 30, 20])[0]))
                    if random.random() > 0.2:  # 80% coverage
                        records.append({"employee": emp, "skill": skill, "proficiency": prof})
            
            clusters = _auto_cluster_skills(all_skills)
            coverage = round(len(set(r["employee"] for r in records)) / max(len(emps), 1) * 100)
            
            validation_errors.append("No Skills tab found — proficiency estimated from work design task skills. Upload a Skills tab for real assessments.")
            
            return _safe({
                "employees": emps,
                "skills": all_skills,
                "clusters": clusters,
                "records": records,
                "coverage": coverage,
                "sample": True,
                "total_headcount": len(emps),
                "validation_errors": validation_errors,
                "dedup_suggestions": [],
            })
    
    # CASE C: No data at all — full sample
    import random
    random.seed(42)
    
    emps = employees[:30] if employees else [f"Employee_{i+1}" for i in range(20)]
    
    skill_clusters = {
        "Technical": ["Data Analysis", "Python/SQL", "Cloud Platforms", "AI/ML Tools", "Process Automation"],
        "Human": ["Communication", "Leadership", "Stakeholder Mgmt", "Negotiation", "Team Collaboration"],
        "Adaptive": ["AI Literacy", "Change Agility", "Critical Thinking", "Creative Problem Solving", "Digital Fluency"],
    }
    all_skills = [s for sks in skill_clusters.values() for s in sks]
    
    records = []
    for emp in emps:
        for skill in all_skills:
            if random.random() > 0.25:
                prof = random.choices([1, 2, 3, 4], weights=[15, 35, 35, 15])[0]
                records.append({"employee": emp, "skill": skill, "proficiency": prof})
    
    validation_errors.append("No workforce or skills data uploaded — showing sample data for preview")
    
    return _safe({
        "employees": emps,
        "skills": all_skills,
        "clusters": skill_clusters,
        "records": records,
        "coverage": round(len(set(r["employee"] for r in records)) / max(len(emps), 1) * 100),
        "sample": True,
        "total_headcount": len(emps),
        "validation_errors": validation_errors,
        "dedup_suggestions": [],
    })


def _auto_cluster_skills(skills: list) -> dict:
    """Auto-cluster skills into Technical/Human/Adaptive."""
    tech_kw = ["data", "python", "sql", "cloud", "api", "automation", "ml", "ai", "engineering", "analytics", "system", "code", "platform", "cyber", "infrastructure", "devops"]
    human_kw = ["communication", "leadership", "stakeholder", "negotiation", "collaboration", "team", "coaching", "mentoring", "presentation", "relationship", "management"]
    adaptive_kw = ["agility", "change", "critical", "creative", "digital", "literacy", "problem", "learning", "innovation", "resilience", "adaptab"]
    
    clusters = {"Technical": [], "Human": [], "Adaptive": []}
    for s in skills:
        sl = s.lower()
        if any(kw in sl for kw in tech_kw):
            clusters["Technical"].append(s)
        elif any(kw in sl for kw in human_kw):
            clusters["Human"].append(s)
        elif any(kw in sl for kw in adaptive_kw):
            clusters["Adaptive"].append(s)
        else:
            clusters["Technical"].append(s)  # default
    
    return _safe({k: v for k, v in clusters.items() if v})


def _find_skill_duplicates(skills: list) -> list:
    """Find potential duplicate skill names."""
    from difflib import SequenceMatcher
    dupes = []
    seen = set()
    for i, a in enumerate(skills):
        for b in skills[i+1:]:
            if (a, b) in seen or (b, a) in seen:
                continue
            ratio = SequenceMatcher(None, a.lower(), b.lower()).ratio()
            if 0.6 < ratio < 1.0:
                dupes.append({"skill_a": a, "skill_b": b, "similarity": round(ratio * 100)})
                seen.add((a, b))
    return sorted(dupes, key=lambda d: -d["similarity"])[:10]


@app.get("/api/skills/gap/{model_id}")
async def get_skills_gap(model_id: str):
    """Skills Gap Analysis — current proficiency vs target from WDL reconstruction."""
    cached = _get_cached("get_skills_gap", model_id)
    if cached is not None:
        return cached
    ds = store.datasets.get(model_id)
    
    # Get current inventory
    inv = await get_skills_inventory(model_id)
    records = inv.get("records", [])
    
    if not records:
        return _safe({"gaps": [], "summary": {"message": "No skills data — complete Skills Inventory first"}, "role_gaps": []})
    
    # Calculate current avg proficiency per skill
    from collections import defaultdict
    skill_data = defaultdict(lambda: {"sum": 0, "count": 0, "employees": []})
    for r in records:
        skill_data[r["skill"]]["sum"] += r["proficiency"]
        skill_data[r["skill"]]["count"] += 1
        skill_data[r["skill"]]["employees"].append({"employee": r["employee"], "proficiency": r["proficiency"]})
    
    # Determine target proficiency
    # Priority 1: From WDL reconstruction (real targets)
    # Priority 2: From task automatability scoring (estimated targets)
    # Priority 3: Default benchmarks
    
    targets = {}
    
    if ds:
        # Try to get targets from WDL reconstruction
        wd_df = None
        for key in ds:
            kl = key.lower()
            if any(t in kl for t in ["work_design", "work design", "tasks", "task", "deconstruction"]):
                wd_df = ds[key]
                break
        
        if wd_df is not None and not wd_df.empty:
            from app.store import sync_work_design, build_reconstruction, build_skill_analysis
            t = sync_work_design(wd_df)
            recon = build_reconstruction(wd_df)
            _, future_skills = build_skill_analysis(wd_df, recon)
            
            if not future_skills.empty:
                # Skills with higher future weight need higher proficiency
                max_weight = future_skills["Weight"].max() if not future_skills.empty else 1
                for _, row in future_skills.iterrows():
                    skill = str(row["Skill"])
                    weight = float(row.get("Weight", 0))
                    # Scale: high-weight future skills need proficiency 3.5-4, low-weight need 2.5-3
                    targets[skill] = round(2.5 + (weight / max(max_weight, 1)) * 1.5, 1)
    
    # AI-era skills get higher default targets
    ai_skills = {"ai", "ml", "automation", "digital", "data", "python", "cloud", "cyber", "analytics"}
    
    gaps = []
    for skill, data in skill_data.items():
        current_avg = round(data["sum"] / data["count"], 1)
        
        # Get target — WDL-derived, or default based on skill type
        if skill in targets:
            target = targets[skill]
            target_source = "Work Design Lab"
        elif any(kw in skill.lower() for kw in ai_skills):
            target = 3.5
            target_source = "AI-era benchmark"
        else:
            target = 3.0
            target_source = "Industry benchmark"
        
        delta = round(target - current_avg, 1)
        
        # Individual employee gaps for drilldown
        emp_gaps = []
        for e in data["employees"]:
            emp_delta = round(target - e["proficiency"], 1)
            emp_gaps.append({
                "employee": e["employee"],
                "current": e["proficiency"],
                "target": target,
                "delta": emp_delta,
                "reskillable": emp_delta <= 2,
            })
        emp_gaps.sort(key=lambda g: -g["delta"])
        
        gaps.append({
            "skill": skill,
            "current_avg": current_avg,
            "target": target,
            "target_source": target_source,
            "delta": max(delta, 0),
            "employees_assessed": data["count"],
            "severity": "Critical" if delta > 1.5 else "Moderate" if delta > 0.5 else "Low",
            "employee_gaps": emp_gaps[:20],
        })
    
    gaps.sort(key=lambda g: -g["delta"])
    
    critical = len([g for g in gaps if g["severity"] == "Critical"])
    moderate = len([g for g in gaps if g["severity"] == "Moderate"])
    low = len([g for g in gaps if g["severity"] == "Low"])
    
    return _safe({
        "gaps": gaps,
        "summary": {
            "total_skills": len(gaps),
            "critical_gaps": critical,
            "moderate_gaps": moderate,
            "low_gaps": low,
            "avg_gap": round(sum(g["delta"] for g in gaps) / max(len(gaps), 1), 1),
            "largest_gap_skill": gaps[0]["skill"] if gaps else "—",
            "largest_gap_delta": gaps[0]["delta"] if gaps else 0,
            "wdl_connected": len(targets) > 0,
        }
    })


@app.get("/api/skills/adjacency/{model_id}")
async def get_skills_adjacency(model_id: str):
    """Skills Adjacency Map — maps current employees to redesigned roles."""
    cached = _get_cached("get_skills_adjacency", model_id)
    if cached is not None:
        return cached
    ds = store.datasets.get(model_id)
    
    inv = await get_skills_inventory(model_id)
    records = inv.get("records", [])
    employees = inv.get("employees", [])
    
    if not records or not employees:
        return _safe({"adjacencies": [], "summary": {}})
    
    # Build employee skill profiles
    from collections import defaultdict
    emp_skills = defaultdict(dict)
    for r in records:
        emp_skills[r["employee"]][r["skill"]] = r["proficiency"]
    
    # Get target roles from WDL reconstruction if available
    target_roles = []
    
    if ds:
        wd_df = None
        for key in ds:
            kl = key.lower()
            if any(t in kl for t in ["work_design", "work design", "tasks", "task"]):
                wd_df = ds[key]
                break
        
        if wd_df is not None and not wd_df.empty:
            from app.store import sync_work_design, build_reconstruction, build_skill_analysis
            recon = build_reconstruction(wd_df)
            _, future_skills = build_skill_analysis(wd_df, recon)
            
            if not future_skills.empty:
                # Group future skills into role-like clusters
                top_skills = future_skills.nlargest(15, "Weight")
                # Create target roles from high-weight skill combinations
                tech_skills = {s: 3 for s in top_skills[top_skills["Skill"].str.lower().str.contains("data|python|sql|cloud|auto|ai|ml|analytics", na=False)]["Skill"].tolist()[:5]}
                human_skills = {s: 3 for s in top_skills[top_skills["Skill"].str.lower().str.contains("communicat|leadership|stakeholder|team|coach|negotiat", na=False)]["Skill"].tolist()[:5]}
                adaptive_skills = {s: 3 for s in top_skills[top_skills["Skill"].str.lower().str.contains("agil|change|critical|creative|digital|innovat", na=False)]["Skill"].tolist()[:5]}
                
                if tech_skills:
                    target_roles.append({"role": "AI Operations Analyst", "required_skills": {**tech_skills, **dict(list(adaptive_skills.items())[:2])}})
                if human_skills and tech_skills:
                    target_roles.append({"role": "AI Change Manager", "required_skills": {**human_skills, **dict(list(adaptive_skills.items())[:2])}})
                if tech_skills:
                    target_roles.append({"role": "Automation Engineer", "required_skills": tech_skills})
    
    # Fallback target roles if WDL hasn't produced any
    if not target_roles:
        all_skills_set = set(r["skill"] for r in records)
        target_roles = [
            {"role": "AI Operations Analyst", "required_skills": {s: 3 for s in list(all_skills_set)[:5]}},
            {"role": "Data Governance Lead", "required_skills": {s: 3 for s in list(all_skills_set)[2:7]}},
            {"role": "Automation Engineer", "required_skills": {s: 3 for s in list(all_skills_set)[4:9]}},
            {"role": "AI Change Manager", "required_skills": {s: 3 for s in list(all_skills_set)[6:11]}},
            {"role": "Digital Transformation Lead", "required_skills": {s: 3 for s in list(all_skills_set)[8:13]}},
        ]
    
    adjacencies = []
    for target in target_roles:
        if not target["required_skills"]:
            continue
        matches = []
        for emp in employees[:50]:
            profile = emp_skills.get(emp, {})
            if not profile:
                continue
            
            match_points = 0
            total_points = 0
            matching_skills = []
            gap_skills = []
            
            for skill, req_level in target["required_skills"].items():
                total_points += req_level
                current = profile.get(skill, 0)
                match_points += min(current, req_level)
                if current >= req_level * 0.7:
                    matching_skills.append(skill)
                elif current < req_level * 0.5:
                    gap_skills.append(skill)
            
            adjacency_pct = round((match_points / max(total_points, 1)) * 100)
            reskill_months = sum(max(0, target["required_skills"].get(s, 0) - profile.get(s, 0)) for s in gap_skills) * 3
            
            matches.append({
                "employee": emp,
                "adjacency_pct": adjacency_pct,
                "matching_skills": matching_skills[:5],
                "gap_skills": gap_skills[:5],
                "reskill_months": reskill_months,
            })
        
        matches.sort(key=lambda m: -m["adjacency_pct"])
        
        adjacencies.append({
            "target_role": target["role"],
            "required_skills": target["required_skills"],
            "top_candidates": matches[:10],
            "strong_matches": len([m for m in matches if m["adjacency_pct"] >= 70]),
            "reskillable": len([m for m in matches if 50 <= m["adjacency_pct"] < 70]),
            "weak_matches": len([m for m in matches if m["adjacency_pct"] < 50]),
            "wdl_derived": len(target.get("required_skills", {})) > 0,
        })
    
    fillable = len([a for a in adjacencies if a["strong_matches"] > 0])
    
    return _safe({
        "adjacencies": adjacencies,
        "summary": {
            "target_roles": len(adjacencies),
            "fillable_internally": fillable,
            "need_external": len(adjacencies) - fillable,
            "avg_best_adjacency": round(sum(a["top_candidates"][0]["adjacency_pct"] for a in adjacencies if a["top_candidates"]) / max(len(adjacencies), 1)) if adjacencies else 0,
            "wdl_connected": any(a.get("wdl_derived") for a in adjacencies),
        }
    })


# ═══════════════════════════════════════════════════
# BUILD/BUY/BORROW/AUTOMATE ENDPOINT
# ═══════════════════════════════════════════════════

@app.get("/api/bbba/{model_id}")
async def get_bbba(model_id: str):
    """Build/Buy/Borrow/Automate framework — per-role disposition with cost implications."""
    cached = _get_cached("get_bbba", model_id)
    if cached is not None:
        return cached
    gap_data = await get_skills_gap(model_id)
    adj_data = await get_skills_adjacency(model_id)
    
    gaps = gap_data.get("gaps", [])
    adjacencies = adj_data.get("adjacencies", [])
    
    if not gaps and not adjacencies:
        return _safe({"roles": [], "summary": {}, "cost_model": {}})
    
    # Build role-level dispositions from gap analysis and adjacency
    roles = []
    for adj in adjacencies:
        role = adj["target_role"]
        strong = adj.get("strong_matches", 0)
        reskillable = adj.get("reskillable", 0)
        
        # Auto-recommend disposition
        if strong >= 3:
            disposition = "Build"
            reason = f"{strong} strong internal candidates available"
            cost_per = 15000  # reskilling cost
        elif reskillable >= 2:
            disposition = "Build"
            reason = f"{reskillable} reskillable candidates with investment"
            cost_per = 25000
        elif strong >= 1:
            disposition = "Build"
            reason = "Limited but viable internal candidates"
            cost_per = 20000
        else:
            disposition = "Buy"
            reason = "No viable internal candidates above threshold"
            cost_per = 85000  # external hire cost
        
        # Check if role can be automated
        role_gaps = [g for g in gaps if g.get("severity") == "Critical"]
        auto_skills = [g for g in role_gaps if "automat" in g.get("skill", "").lower() or "ai" in g.get("skill", "").lower()]
        if len(auto_skills) > len(role_gaps) * 0.5 and len(role_gaps) > 0:
            disposition = "Automate"
            reason = "Majority of skill needs can be met by AI/automation"
            cost_per = 50000  # one-time automation cost
        
        # Estimate FTE need
        fte_need = max(1, strong + reskillable) if disposition == "Build" else 2
        
        roles.append({
            "role": role,
            "disposition": disposition,
            "reason": reason,
            "strong_candidates": strong,
            "reskillable_candidates": reskillable,
            "cost_per_fte": cost_per,
            "fte_needed": fte_need,
            "total_cost": cost_per * fte_need,
            "required_skills": list(adj.get("required_skills", {}).keys()),
            "timeline_months": 3 if disposition == "Build" and strong >= 2 else 6 if disposition == "Build" else 2 if disposition == "Buy" else 4,
        })
    
    # Add generic roles for gaps without adjacency matches
    for g in gaps:
        if g["severity"] == "Critical" and not any(g["skill"] in r.get("required_skills", []) for r in roles):
            roles.append({
                "role": f"{g['skill']} Specialist",
                "disposition": "Buy" if g["delta"] > 2 else "Borrow",
                "reason": f"Critical gap (delta {g['delta']}) — {'too large to reskill' if g['delta'] > 2 else 'contract to bridge'}",
                "strong_candidates": 0,
                "reskillable_candidates": 0,
                "cost_per_fte": 95000 if g["delta"] > 2 else 45000,
                "fte_needed": 1,
                "total_cost": 95000 if g["delta"] > 2 else 45000,
                "required_skills": [g["skill"]],
                "timeline_months": 2,
            })
    
    # Summary
    build_count = len([r for r in roles if r["disposition"] == "Build"])
    buy_count = len([r for r in roles if r["disposition"] == "Buy"])
    borrow_count = len([r for r in roles if r["disposition"] == "Borrow"])
    auto_count = len([r for r in roles if r["disposition"] == "Automate"])
    
    return _safe({
        "roles": roles,
        "summary": {
            "total_roles": len(roles),
            "build": build_count,
            "buy": buy_count,
            "borrow": borrow_count,
            "automate": auto_count,
            "total_investment": sum(r["total_cost"] for r in roles),
            "reskilling_investment": sum(r["total_cost"] for r in roles if r["disposition"] == "Build"),
            "hiring_cost": sum(r["total_cost"] for r in roles if r["disposition"] == "Buy"),
        },
        "cost_model": {
            "build_avg": 20000,
            "buy_avg": 85000,
            "borrow_avg": 45000,
            "automate_avg": 50000,
        }
    })


# ═══════════════════════════════════════════════════
# HEADCOUNT PLANNING ENDPOINT
# ═══════════════════════════════════════════════════

@app.get("/api/headcount/{model_id}")
async def get_headcount_plan(model_id: str):
    """Headcount planning — current vs future FTE with waterfall."""
    cached = _get_cached("get_headcount_plan", model_id)
    if cached is not None:
        return cached
    ds = store.datasets.get(model_id)
    bbba = await get_bbba(model_id)
    inv = await get_skills_inventory(model_id)
    
    total_current = len(inv.get("employees", []))
    roles = bbba.get("roles", [])
    
    # Calculate waterfall
    build_roles = [r for r in roles if r["disposition"] == "Build"]
    buy_roles = [r for r in roles if r["disposition"] == "Buy"]
    borrow_roles = [r for r in roles if r["disposition"] == "Borrow"]
    auto_roles = [r for r in roles if r["disposition"] == "Automate"]
    
    eliminated = sum(r["fte_needed"] for r in auto_roles)
    redeployed = sum(r["fte_needed"] for r in build_roles)
    new_hires = sum(r["fte_needed"] for r in buy_roles)
    contractors = sum(r["fte_needed"] for r in borrow_roles)
    
    # Estimate attrition (industry avg ~12% annual, 18mo horizon = ~18%)
    attrition_rate = 0.18
    natural_attrition = round(total_current * attrition_rate)
    
    target = total_current - eliminated + new_hires - min(natural_attrition, eliminated)
    
    # Department breakdown (estimate)
    import random
    random.seed(hash(model_id) % 2**31)
    depts = ["Finance", "Technology", "Operations", "HR", "Marketing", "Legal"]
    dept_data = []
    remaining = total_current
    for i, dept in enumerate(depts):
        if i == len(depts) - 1:
            current = remaining
        else:
            current = max(1, round(total_current * random.uniform(0.1, 0.25)))
            remaining -= current
        
        dept_eliminated = round(current * random.uniform(0, 0.15))
        dept_new = round(current * random.uniform(0, 0.1))
        
        dept_data.append({
            "department": dept,
            "current_fte": current,
            "eliminated": dept_eliminated,
            "redeployed": round(current * random.uniform(0.05, 0.15)),
            "new_hires": dept_new,
            "future_fte": current - dept_eliminated + dept_new,
            "pct_change": round(((current - dept_eliminated + dept_new) / max(current, 1) - 1) * 100, 1),
        })
    
    return _safe({
        "waterfall": {
            "starting_headcount": total_current,
            "eliminations": eliminated,
            "natural_attrition": min(natural_attrition, eliminated),
            "redeployments": redeployed,
            "new_hires": new_hires,
            "contractors": contractors,
            "target_headcount": target,
            "net_change": target - total_current,
            "net_change_pct": round(((target / max(total_current, 1)) - 1) * 100, 1),
        },
        "departments": dept_data,
        "timeline": {
            "phase_1_months": "1-6",
            "phase_1_actions": f"Redeploy {redeployed} FTE, begin reskilling",
            "phase_2_months": "7-12",
            "phase_2_actions": f"Hire {new_hires} new roles, onboard {contractors} contractors",
            "phase_3_months": "13-18",
            "phase_3_actions": f"Complete transitions, absorb {min(natural_attrition, eliminated)} through attrition",
        },
    })


# ═══════════════════════════════════════════════════
# AI READINESS ASSESSMENT ENDPOINT
# ═══════════════════════════════════════════════════

@app.get("/api/readiness/{model_id}")
async def get_readiness_assessment(model_id: str):
    """AI Readiness — individual, team, and org-level scoring."""
    inv = await get_skills_inventory(model_id)
    employees = inv.get("employees", [])
    records = inv.get("records", [])
    
    import random
    random.seed(hash(model_id) % 2**31)
    
    dimensions = ["AI Awareness", "Tool Adoption", "Data Literacy", "Change Openness", "AI Collaboration"]
    
    # Generate individual scores
    individuals = []
    for emp in employees[:50]:
        scores = {}
        # Base scores from skill proficiency if available
        emp_recs = [r for r in records if r["employee"] == emp]
        base = sum(r["proficiency"] for r in emp_recs) / max(len(emp_recs), 1) if emp_recs else 2
        
        for dim in dimensions:
            scores[dim] = round(max(1, min(5, base + random.uniform(-1.2, 1.2))), 1)
        
        avg = round(sum(scores.values()) / len(scores), 1)
        band = "Ready Now" if avg >= 3.8 else "Coachable" if avg >= 2.5 else "At Risk"
        
        individuals.append({
            "employee": emp,
            "scores": scores,
            "average": avg,
            "band": band,
        })
    
    # Team/function rollup
    ready = len([i for i in individuals if i["band"] == "Ready Now"])
    coachable = len([i for i in individuals if i["band"] == "Coachable"])
    at_risk = len([i for i in individuals if i["band"] == "At Risk"])
    
    # Dimension averages
    dim_avgs = {}
    for dim in dimensions:
        vals = [i["scores"][dim] for i in individuals]
        dim_avgs[dim] = round(sum(vals) / max(len(vals), 1), 1)
    
    return _safe({
        "individuals": individuals,
        "dimensions": dimensions,
        "dimension_averages": dim_avgs,
        "bands": {"ready_now": ready, "coachable": coachable, "at_risk": at_risk},
        "org_average": round(sum(i["average"] for i in individuals) / max(len(individuals), 1), 1),
        "lowest_dimension": min(dim_avgs, key=dim_avgs.get) if dim_avgs else "—",
        "highest_dimension": max(dim_avgs, key=dim_avgs.get) if dim_avgs else "—",
    })


# ═══════════════════════════════════════════════════
# RESKILLING PATHWAYS ENDPOINT
# ═══════════════════════════════════════════════════

@app.get("/api/reskilling/{model_id}")
async def get_reskilling_pathways(model_id: str):
    """Reskilling pathways — per-employee learning plans."""
    cached = _get_cached("get_reskilling_pathways", model_id)
    if cached is not None:
        return cached
    gap_data = await get_skills_gap(model_id)
    bbba_data = await get_bbba(model_id)
    readiness = await get_readiness_assessment(model_id)
    
    gaps = gap_data.get("gaps", [])
    roles = bbba_data.get("roles", [])
    individuals = readiness.get("individuals", [])
    
    build_roles = [r for r in roles if r["disposition"] == "Build"]
    
    pathways = []
    for role in build_roles:
        # Find employees that could fill this role
        role_skills = role.get("required_skills", [])
        
        for ind in individuals[:20]:
            emp = ind["employee"]
            readiness_score = ind["average"]
            
            # Calculate which skills this employee needs
            emp_gaps = []
            for g in gaps:
                if g["skill"] in role_skills:
                    emp_gap_data = [eg for eg in g.get("employee_gaps", []) if eg["employee"] == emp]
                    if emp_gap_data and emp_gap_data[0]["delta"] > 0:
                        emp_gaps.append({
                            "skill": g["skill"],
                            "current": emp_gap_data[0]["current"],
                            "target": emp_gap_data[0]["target"],
                            "delta": emp_gap_data[0]["delta"],
                            "intervention": "Course" if emp_gap_data[0]["delta"] <= 1 else "Coaching + Course" if emp_gap_data[0]["delta"] <= 2 else "Intensive Program",
                            "months": round(emp_gap_data[0]["delta"] * 3),
                        })
            
            if emp_gaps:
                total_months = max(eg["months"] for eg in emp_gaps) if emp_gaps else 0
                pathways.append({
                    "employee": emp,
                    "target_role": role["role"],
                    "readiness_score": readiness_score,
                    "readiness_band": ind["band"],
                    "skills_to_develop": emp_gaps,
                    "total_months": total_months,
                    "estimated_cost": sum(5000 if eg["intervention"] == "Course" else 12000 if eg["intervention"] == "Coaching + Course" else 25000 for eg in emp_gaps),
                    "priority": "High" if readiness_score >= 3.5 and total_months <= 6 else "Medium" if readiness_score >= 2.5 else "Low",
                })
    
    pathways.sort(key=lambda p: ({"High": 0, "Medium": 1, "Low": 2}[p["priority"]], p["total_months"]))
    
    return _safe({
        "pathways": pathways[:30],
        "summary": {
            "total_employees": len(pathways),
            "high_priority": len([p for p in pathways if p["priority"] == "High"]),
            "medium_priority": len([p for p in pathways if p["priority"] == "Medium"]),
            "avg_months": round(sum(p["total_months"] for p in pathways) / max(len(pathways), 1), 1),
            "total_investment": sum(p["estimated_cost"] for p in pathways),
        }
    })


# ═══════════════════════════════════════════════════
# TALENT MARKETPLACE ENDPOINT
# ═══════════════════════════════════════════════════

@app.get("/api/marketplace/{model_id}")
async def get_talent_marketplace(model_id: str):
    """Talent marketplace — ranked internal candidates per redesigned role."""
    cached = _get_cached("get_talent_marketplace", model_id)
    if cached is not None:
        return cached
    adj_data = await get_skills_adjacency(model_id)
    readiness = await get_readiness_assessment(model_id)
    reskilling = await get_reskilling_pathways(model_id)
    
    adjacencies = adj_data.get("adjacencies", [])
    individuals = {i["employee"]: i for i in readiness.get("individuals", [])}
    pathways_by_emp = {}
    for p in reskilling.get("pathways", []):
        pathways_by_emp.setdefault(p["employee"], []).append(p)
    
    marketplace = []
    for adj in adjacencies:
        role_matches = []
        for cand in adj.get("top_candidates", []):
            emp = cand["employee"]
            ind = individuals.get(emp, {})
            paths = pathways_by_emp.get(emp, [])
            
            role_matches.append({
                "employee": emp,
                "adjacency_pct": cand["adjacency_pct"],
                "matching_skills": cand.get("matching_skills", []),
                "gap_skills": cand.get("gap_skills", []),
                "reskill_months": cand.get("reskill_months", 0),
                "readiness_score": ind.get("average", 0),
                "readiness_band": ind.get("band", "—"),
                "has_pathway": len(paths) > 0,
                "pathway_cost": paths[0]["estimated_cost"] if paths else 0,
                "composite_score": round(cand["adjacency_pct"] * 0.5 + ind.get("average", 0) * 10 * 0.3 + (100 - cand.get("reskill_months", 6) * 5) * 0.2),
            })
        
        role_matches.sort(key=lambda m: -m["composite_score"])
        
        marketplace.append({
            "target_role": adj["target_role"],
            "candidates": role_matches[:8],
            "fill_recommendation": "Internal" if role_matches and role_matches[0]["adjacency_pct"] >= 65 else "External",
        })
    
    return _safe({
        "marketplace": marketplace,
        "summary": {
            "total_roles": len(marketplace),
            "internal_fill": len([m for m in marketplace if m["fill_recommendation"] == "Internal"]),
            "external_fill": len([m for m in marketplace if m["fill_recommendation"] == "External"]),
        }
    })


# ═══════════════════════════════════════════════════
# MANAGER CAPABILITY ASSESSMENT
# ═══════════════════════════════════════════════════

@app.get("/api/manager-capability/{model_id}")
async def get_manager_capability(model_id: str):
    """Manager capability assessment — score managers, correlate with team readiness."""
    readiness = await get_readiness_assessment(model_id)
    inv = await get_skills_inventory(model_id)
    
    individuals = readiness.get("individuals", [])
    employees = inv.get("employees", [])
    
    import random
    random.seed(hash(model_id + "mgr") % 2**31)
    
    dimensions = ["Leading Through Ambiguity", "AI Tool Fluency", "Coaching & Reskilling", "Communication of Change"]
    
    # Generate managers (people with direct reports)
    num_managers = max(3, len(employees) // 8)
    manager_pool = employees[:num_managers] if employees else [f"Manager_{i+1}" for i in range(8)]
    
    managers = []
    for mgr in manager_pool:
        scores = {}
        for dim in dimensions:
            scores[dim] = round(max(1, min(5, 2.5 + random.gauss(0, 1))), 1)
        
        avg = round(sum(scores.values()) / len(scores), 1)
        category = "Transformation Champion" if avg >= 3.8 else "Needs Development" if avg >= 2.5 else "Flight Risk"
        
        # Get team members and their readiness
        team_size = random.randint(3, 12)
        team_members = random.sample(individuals, min(team_size, len(individuals))) if individuals else []
        team_readiness_avg = round(sum(m["average"] for m in team_members) / max(len(team_members), 1), 1) if team_members else 0
        
        managers.append({
            "manager": mgr,
            "scores": scores,
            "average": avg,
            "category": category,
            "direct_reports": team_size,
            "team_readiness_avg": team_readiness_avg,
            "correlation": round(min(1, max(-1, (avg - 3) * 0.4 + (team_readiness_avg - 3) * 0.3 + random.uniform(-0.2, 0.2))), 2),
            "team_members": [{"employee": m["employee"], "readiness": m["average"], "band": m["band"]} for m in team_members[:5]],
        })
    
    champions = len([m for m in managers if m["category"] == "Transformation Champion"])
    needs_dev = len([m for m in managers if m["category"] == "Needs Development"])
    flight_risk = len([m for m in managers if m["category"] == "Flight Risk"])
    
    # Correlation insight
    high_mgr_teams = [m for m in managers if m["average"] >= 3.5]
    low_mgr_teams = [m for m in managers if m["average"] < 2.5]
    high_team_avg = round(sum(m["team_readiness_avg"] for m in high_mgr_teams) / max(len(high_mgr_teams), 1), 1) if high_mgr_teams else 0
    low_team_avg = round(sum(m["team_readiness_avg"] for m in low_mgr_teams) / max(len(low_mgr_teams), 1), 1) if low_mgr_teams else 0
    
    return _safe({
        "managers": managers,
        "dimensions": dimensions,
        "summary": {
            "total_managers": len(managers),
            "champions": champions,
            "needs_development": needs_dev,
            "flight_risk": flight_risk,
            "avg_score": round(sum(m["average"] for m in managers) / max(len(managers), 1), 1),
            "weakest_dimension": min(dimensions, key=lambda d: sum(m["scores"][d] for m in managers) / max(len(managers), 1)) if managers else "—",
            "correlation_multiplier": f"{round(high_team_avg / max(low_team_avg, 0.1), 1)}x" if low_team_avg > 0 else "—",
            "high_mgr_team_readiness": high_team_avg,
            "low_mgr_team_readiness": low_team_avg,
        }
    })


# ═══════════════════════════════════════════════════
# CHANGE READINESS & ADOPTION PROPENSITY
# ═══════════════════════════════════════════════════

@app.get("/api/change-readiness/{model_id}")
async def get_change_readiness(model_id: str):
    """Change readiness — 4-quadrant segmentation with intervention mapping."""
    readiness = await get_readiness_assessment(model_id)
    gap_data = await get_skills_gap(model_id)
    
    individuals = readiness.get("individuals", [])
    gaps = gap_data.get("gaps", [])
    
    import random
    random.seed(hash(model_id + "chg") % 2**31)
    
    # Calculate impact score per employee (based on how many critical gaps affect their role)
    avg_gap = sum(g["delta"] for g in gaps) / max(len(gaps), 1)
    
    quadrants = {"high_ready_high_impact": [], "high_ready_low_impact": [], "low_ready_high_impact": [], "low_ready_low_impact": []}
    
    for ind in individuals:
        readiness_score = ind["average"]
        # Simulate impact score — in real tool this comes from WDL task analysis per role
        impact_score = round(max(1, min(5, avg_gap * 1.2 + random.uniform(-1, 1))), 1)
        
        is_high_ready = readiness_score >= 3.0
        is_high_impact = impact_score >= 3.0
        
        entry = {
            "employee": ind["employee"],
            "readiness": readiness_score,
            "impact": impact_score,
            "band": ind["band"],
        }
        
        if is_high_ready and is_high_impact:
            quadrants["high_ready_high_impact"].append(entry)
        elif is_high_ready and not is_high_impact:
            quadrants["high_ready_low_impact"].append(entry)
        elif not is_high_ready and is_high_impact:
            quadrants["low_ready_high_impact"].append(entry)
        else:
            quadrants["low_ready_low_impact"].append(entry)
    
    # Intervention recommendations per quadrant
    interventions = {
        "high_ready_high_impact": {"label": "Champions", "color": "#10B981", "action": "Leverage as early adopters and peer champions", "cadence": "Monthly check-in", "priority": 2},
        "high_ready_low_impact": {"label": "Supporters", "color": "#3B82F6", "action": "Light-touch communication, keep informed", "cadence": "Quarterly update", "priority": 4},
        "low_ready_high_impact": {"label": "High Risk", "color": "#EF4444", "action": "Dedicated change champion, bi-weekly touchpoints, intensive support", "cadence": "Bi-weekly 1:1", "priority": 1},
        "low_ready_low_impact": {"label": "Monitor", "color": "#F59E0B", "action": "Standard communications, group sessions", "cadence": "Monthly town hall", "priority": 3},
    }
    
    # Function concentration
    total = len(individuals)
    high_risk_count = len(quadrants["low_ready_high_impact"])
    
    return _safe({
        "quadrants": {k: {"employees": v, "count": len(v), **interventions[k]} for k, v in quadrants.items()},
        "summary": {
            "total_assessed": total,
            "high_risk_count": high_risk_count,
            "high_risk_pct": round(high_risk_count / max(total, 1) * 100),
            "champion_count": len(quadrants["high_ready_high_impact"]),
            "recommended_champions_needed": max(1, high_risk_count // 5),
        },
        "messaging_guidance": {
            "high_risk": "Focus on: what's changing, what's NOT changing, personal support available, timeline, and quick wins they'll see early",
            "champions": "Focus on: their role in helping peers, early access to tools, recognition, feedback channels",
            "supporters": "Focus on: keeping informed, answering questions proactively, showing progress",
            "monitor": "Focus on: general awareness, optional deep-dive sessions, self-service resources",
        }
    })


# ═══════════════════════════════════════════════════
# MANAGER DEVELOPMENT TRACK
# ═══════════════════════════════════════════════════

@app.get("/api/manager-development/{model_id}")
async def get_manager_development(model_id: str):
    """Manager development — targeted plans per capability band."""
    mgr_data = await get_manager_capability(model_id)
    
    managers = mgr_data.get("managers", [])
    dimensions = mgr_data.get("dimensions", [])
    
    tracks = []
    for mgr in managers:
        weak_dims = sorted(dimensions, key=lambda d: mgr["scores"][d])[:2]
        
        interventions = []
        for dim in weak_dims:
            score = mgr["scores"][dim]
            if score < 2:
                interventions.append({"dimension": dim, "intervention": "Intensive Coaching Program", "format": "1:1 Executive Coach", "duration_weeks": 12, "cost": 8000})
            elif score < 3:
                interventions.append({"dimension": dim, "intervention": "Development Workshop Series", "format": "Group + 1:1", "duration_weeks": 8, "cost": 4000})
            else:
                interventions.append({"dimension": dim, "intervention": "Peer Learning Circle", "format": "Group Sessions", "duration_weeks": 6, "cost": 1500})
        
        role_in_change = "Change Agent — lead transformation for their team" if mgr["category"] == "Transformation Champion" else "Development Priority — build capability before rollout" if mgr["category"] == "Needs Development" else "Retention Risk — engage immediately, assess commitment"
        
        tracks.append({
            "manager": mgr["manager"],
            "category": mgr["category"],
            "average_score": mgr["average"],
            "direct_reports": mgr["direct_reports"],
            "weak_dimensions": weak_dims,
            "interventions": interventions,
            "role_in_change": role_in_change,
            "total_cost": sum(i["cost"] for i in interventions),
            "total_weeks": max(i["duration_weeks"] for i in interventions) if interventions else 0,
        })
    
    tracks.sort(key=lambda t: {"Transformation Champion": 2, "Needs Development": 1, "Flight Risk": 0}[t["category"]])
    
    return _safe({
        "tracks": tracks,
        "summary": {
            "total_managers": len(tracks),
            "total_investment": sum(t["total_cost"] for t in tracks),
            "avg_duration_weeks": round(sum(t["total_weeks"] for t in tracks) / max(len(tracks), 1), 1),
            "change_agents": len([t for t in tracks if t["category"] == "Transformation Champion"]),
            "need_development": len([t for t in tracks if t["category"] == "Needs Development"]),
            "flight_risks": len([t for t in tracks if t["category"] == "Flight Risk"]),
        }
    })


# ═══════════════════════════════════════════════════
# EXPORT / REPORT GENERATOR
# ═══════════════════════════════════════════════════

@app.get("/api/export/summary/{model_id}")
async def get_export_summary(model_id: str):
    """Gather all module data for report generation."""
    inv = await get_skills_inventory(model_id)
    gap = await get_skills_gap(model_id)
    adj = await get_skills_adjacency(model_id)
    bbba = await get_bbba(model_id)
    hc = await get_headcount_plan(model_id)
    readiness = await get_readiness_assessment(model_id)
    mgr = await get_manager_capability(model_id)
    chg = await get_change_readiness(model_id)
    reskill = await get_reskilling_pathways(model_id)
    mp = await get_talent_marketplace(model_id)
    
    return _safe({
        "skills_coverage": inv.get("coverage", 0),
        "total_employees": len(inv.get("employees", [])),
        "critical_gaps": gap.get("summary", {}).get("critical_gaps", 0),
        "fillable_internally": adj.get("summary", {}).get("fillable_internally", 0),
        "bbba_summary": bbba.get("summary", {}),
        "headcount_waterfall": hc.get("waterfall", {}),
        "org_readiness": readiness.get("org_average", 0),
        "readiness_bands": readiness.get("bands", {}),
        "manager_summary": mgr.get("summary", {}),
        "high_risk_pct": chg.get("summary", {}).get("high_risk_pct", 0),
        "reskilling_summary": reskill.get("summary", {}),
        "marketplace_summary": mp.get("summary", {}),
    })


# ═══════════════════════════════════════════════════
# TUTORIAL DATA ENDPOINT
# ═══════════════════════════════════════════════════

TUTORIAL_WORKFORCE = [
    {"Employee ID": f"EMP{str(i+1).zfill(3)}", "Name": name, "Job Title": role, "Function": func, "Level": level, "Manager ID": mgr, "Compensation": comp}
    for i, (name, role, func, level, mgr, comp) in enumerate([
        ("Sarah Chen", "Data Analyst", "Technology", "L3", "EMP015", 85000),
        ("James Rodriguez", "Senior Data Analyst", "Technology", "L4", "EMP015", 105000),
        ("Maria Thompson", "Financial Analyst", "Finance", "L3", "EMP016", 90000),
        ("David Kim", "Operations Coordinator", "Operations", "L2", "EMP017", 65000),
        ("Lisa Patel", "HR Business Partner", "HR", "L4", "EMP018", 95000),
        ("Michael Brown", "Marketing Manager", "Marketing", "L4", "EMP018", 110000),
        ("Emma Wilson", "Software Engineer", "Technology", "L3", "EMP015", 120000),
        ("Robert Taylor", "Process Improvement Spec", "Operations", "L3", "EMP017", 80000),
        ("Jennifer Garcia", "Compliance Analyst", "Legal", "L3", "EMP019", 88000),
        ("William Davis", "Product Manager", "Product", "L4", "EMP015", 130000),
        ("Amanda Martinez", "UX Designer", "Product", "L3", "EMP015", 95000),
        ("Christopher Lee", "Account Executive", "Marketing", "L3", "EMP018", 85000),
        ("Nicole Johnson", "Talent Acquisition Lead", "HR", "L3", "EMP018", 78000),
        ("Daniel White", "Supply Chain Analyst", "Operations", "L3", "EMP017", 75000),
        ("Rachel Anderson", "VP Technology", "Technology", "L5", "EMP020", 185000),
        ("Thomas Harris", "VP Finance", "Finance", "L5", "EMP020", 175000),
        ("Catherine Clark", "VP Operations", "Operations", "L5", "EMP020", 165000),
        ("Steven Wright", "VP HR", "HR", "L5", "EMP020", 155000),
        ("Laura King", "Director Legal", "Legal", "L5", "EMP020", 160000),
        ("Andrew Scott", "CEO", "Executive", "L6", "", 250000),
    ])
]

@app.get("/api/tutorial/check/{project_id}")
async def check_tutorial(project_id: str):
    """Check if this is a tutorial project."""
    return {"is_tutorial": project_id == "tutorial_project"}


# ═══════════════════════════════════════════════════
# TUTORIAL SANDBOX — Auto-seeds rich dataset
# ═══════════════════════════════════════════════════

def _generate_sandbox_dataset(industry: str, size: str):
    """Generate a complete, internally consistent dataset for any industry × size combination."""
    import pandas as pd
    import numpy as np
    import random
    
    seed_val = hash(f"{industry}_{size}") % 2**31
    random.seed(seed_val)
    np.random.seed(seed_val % 2**31)
    
    # ═══ SIZE CONFIG ═══
    size_config = {
        "small":  {"employees": 50,    "functions": 4, "levels": 4, "comp_mult": 0.85, "label": "Small (50)"},
        "medium": {"employees": 200,   "functions": 6, "levels": 5, "comp_mult": 1.0,  "label": "Medium (200)"},
        "large":  {"employees": 1000,  "functions": 8, "levels": 6, "comp_mult": 1.15, "label": "Large (1,000)"},
        "mega":   {"employees": 5000,  "functions": 10, "levels": 7, "comp_mult": 1.35, "label": "Mega (5,000)"},
    }
    sc = size_config.get(size, size_config["medium"])
    num_emp = sc["employees"]
    
    # ═══ INDUSTRY CONFIG ═══
    industry_configs = {
        "financial_services": {
            "label": "Financial Services",
            "company": "Atlas Capital Group",
            "functions": {
                "Trading & Markets": {"roles": ["Quantitative Analyst", "Trader", "Risk Analyst", "Market Data Specialist", "Derivatives Analyst"], "comp_base": 140000},
                "Investment Banking": {"roles": ["Investment Analyst", "Associate", "VP Banking", "Deal Originator", "Pitch Book Specialist"], "comp_base": 130000},
                "Risk & Compliance": {"roles": ["Compliance Officer", "Regulatory Analyst", "AML Analyst", "Audit Manager", "Risk Modeler"], "comp_base": 110000},
                "Wealth Management": {"roles": ["Financial Advisor", "Portfolio Manager", "Client Associate", "Wealth Planner", "Trust Officer"], "comp_base": 120000},
                "Technology": {"roles": ["Software Engineer", "Data Engineer", "Cloud Architect", "InfoSec Analyst", "Platform Engineer"], "comp_base": 135000},
                "Operations": {"roles": ["Operations Analyst", "Settlement Specialist", "Reconciliation Analyst", "Trade Support", "Middle Office Analyst"], "comp_base": 85000},
                "Finance": {"roles": ["Financial Controller", "FP&A Analyst", "Treasury Analyst", "Tax Specialist", "Accountant"], "comp_base": 100000},
                "HR": {"roles": ["HR Business Partner", "Talent Acquisition", "Compensation Analyst", "L&D Specialist", "HRIS Analyst"], "comp_base": 90000},
                "Legal": {"roles": ["General Counsel", "Regulatory Attorney", "Contracts Manager", "Legal Analyst", "Paralegal"], "comp_base": 115000},
                "Marketing": {"roles": ["Brand Manager", "Digital Marketing Lead", "Content Strategist", "Product Marketing", "Events Manager"], "comp_base": 95000},
            },
            "tasks_by_role": {
                "Quantitative Analyst": [("Model development & backtesting", 30, "Variable", "Probabilistic", "Interactive", "Moderate", "Quantitative Analysis", "Python/R"),
                    ("Data pipeline management", 20, "Repetitive", "Deterministic", "Independent", "High", "Data Engineering", "Python/SQL"),
                    ("Research & signal generation", 20, "Variable", "Judgment-heavy", "Independent", "Low", "Financial Markets", "Critical Thinking"),
                    ("Strategy performance reporting", 15, "Repetitive", "Deterministic", "Independent", "High", "Data Analysis", "Communication"),
                    ("Stakeholder presentation", 15, "Variable", "Judgment-heavy", "Collaborative", "Low", "Communication", "Stakeholder Mgmt")],
                "Compliance Officer": [("Regulatory monitoring & updates", 25, "Variable", "Probabilistic", "Interactive", "Moderate", "Regulatory Knowledge", "Research"),
                    ("Policy documentation", 20, "Repetitive", "Deterministic", "Independent", "High", "Technical Writing", "Compliance"),
                    ("Surveillance & monitoring", 20, "Repetitive", "Deterministic", "Independent", "High", "Data Analysis", "Risk Assessment"),
                    ("Training & advisory", 20, "Variable", "Judgment-heavy", "Collaborative", "Low", "Communication", "Leadership"),
                    ("Audit preparation", 15, "Variable", "Probabilistic", "Interactive", "Moderate", "Documentation", "Process Management")],
            },
            "skills": ["Quantitative Analysis", "Python/R", "Financial Markets", "Risk Assessment", "Regulatory Knowledge", "Data Analysis", "SQL/Databases", "Communication", "Stakeholder Mgmt", "Leadership", "AI/ML Tools", "Cloud Platforms", "Process Automation", "Critical Thinking", "Change Agility"],
        },
        "technology": {
            "label": "Technology",
            "company": "Nexus Technologies Inc",
            "functions": {
                "Engineering": {"roles": ["Software Engineer", "Senior Engineer", "Staff Engineer", "Frontend Developer", "Backend Developer"], "comp_base": 145000},
                "Product": {"roles": ["Product Manager", "Senior PM", "Product Analyst", "UX Researcher", "Technical Writer"], "comp_base": 130000},
                "Design": {"roles": ["UX Designer", "UI Designer", "Design Lead", "Design Systems Engineer", "UX Writer"], "comp_base": 120000},
                "Data Science": {"roles": ["Data Scientist", "ML Engineer", "Data Analyst", "Analytics Engineer", "AI Researcher"], "comp_base": 150000},
                "DevOps / SRE": {"roles": ["SRE Engineer", "DevOps Engineer", "Platform Engineer", "Cloud Architect", "Security Engineer"], "comp_base": 155000},
                "Sales": {"roles": ["Account Executive", "Sales Engineer", "SDR", "Enterprise Sales", "Solutions Architect"], "comp_base": 120000},
                "Marketing": {"roles": ["Growth Manager", "Content Marketer", "Product Marketing", "Demand Gen", "Brand Manager"], "comp_base": 110000},
                "People Ops": {"roles": ["People Partner", "Recruiter", "People Analyst", "DEI Lead", "Total Rewards"], "comp_base": 100000},
                "Finance": {"roles": ["FP&A Analyst", "Controller", "Revenue Operations", "Billing Specialist", "Tax Analyst"], "comp_base": 105000},
                "Legal": {"roles": ["Legal Counsel", "Privacy Counsel", "IP Attorney", "Contracts Manager", "Legal Ops"], "comp_base": 130000},
            },
            "tasks_by_role": {
                "Software Engineer": [("Feature development & coding", 35, "Variable", "Probabilistic", "Interactive", "Moderate", "Software Engineering", "Programming"),
                    ("Code review & PR management", 15, "Variable", "Judgment-heavy", "Collaborative", "Low", "Code Quality", "Communication"),
                    ("Testing & QA", 15, "Repetitive", "Deterministic", "Independent", "High", "Testing", "Automation"),
                    ("Documentation", 10, "Repetitive", "Deterministic", "Independent", "High", "Technical Writing", "Communication"),
                    ("Sprint planning & standups", 10, "Variable", "Judgment-heavy", "Collaborative", "Low", "Project Management", "Communication"),
                    ("Incident response & debugging", 15, "Variable", "Probabilistic", "Interactive", "Moderate", "Debugging", "System Design")],
                "Data Scientist": [("Model training & experimentation", 30, "Variable", "Probabilistic", "Independent", "Moderate", "ML/AI", "Python"),
                    ("Data wrangling & feature engineering", 25, "Repetitive", "Deterministic", "Independent", "High", "Data Engineering", "SQL"),
                    ("Stakeholder consultation", 15, "Variable", "Judgment-heavy", "Collaborative", "Low", "Communication", "Business Acumen"),
                    ("Model deployment & monitoring", 15, "Variable", "Probabilistic", "Interactive", "Moderate", "MLOps", "Cloud Platforms"),
                    ("Research & prototyping", 15, "Variable", "Judgment-heavy", "Independent", "Low", "Research", "Innovation")],
            },
            "skills": ["Software Engineering", "Python", "JavaScript/TypeScript", "Cloud Platforms", "ML/AI", "Data Engineering", "System Design", "Product Thinking", "Communication", "Leadership", "Agile/Scrum", "DevOps", "Security", "Critical Thinking", "Innovation"],
        },
        "healthcare": {
            "label": "Healthcare",
            "company": "Meridian Health Systems",
            "functions": {
                "Clinical Operations": {"roles": ["Nurse Manager", "Clinical Coordinator", "Patient Care Specialist", "Medical Technologist", "Pharmacy Tech"], "comp_base": 80000},
                "Medical Staff": {"roles": ["Physician", "Specialist", "Resident", "Physician Assistant", "Nurse Practitioner"], "comp_base": 200000},
                "Administration": {"roles": ["Hospital Administrator", "Department Director", "Operations Manager", "Quality Manager", "Credentialing Specialist"], "comp_base": 95000},
                "Health IT": {"roles": ["EHR Analyst", "Clinical Informaticist", "Health Data Analyst", "Integration Engineer", "Cybersecurity Analyst"], "comp_base": 100000},
                "Revenue Cycle": {"roles": ["Medical Coder", "Billing Specialist", "Claims Analyst", "Authorization Coordinator", "Revenue Integrity Analyst"], "comp_base": 65000},
                "Quality & Compliance": {"roles": ["Compliance Officer", "Quality Improvement Specialist", "Risk Manager", "Accreditation Coordinator", "Patient Safety Officer"], "comp_base": 90000},
                "Finance": {"roles": ["Controller", "Financial Analyst", "Budget Manager", "Reimbursement Analyst", "Payroll Manager"], "comp_base": 85000},
                "HR": {"roles": ["HR Director", "Recruiter", "Benefits Coordinator", "Workforce Planner", "Employee Relations"], "comp_base": 80000},
                "Supply Chain": {"roles": ["Supply Chain Manager", "Procurement Specialist", "Inventory Analyst", "Logistics Coordinator", "Vendor Manager"], "comp_base": 75000},
                "Facilities": {"roles": ["Facilities Director", "Maintenance Supervisor", "Biomedical Engineer", "Safety Officer", "Environmental Services"], "comp_base": 70000},
            },
            "tasks_by_role": {},
            "skills": ["Clinical Knowledge", "EHR Systems", "Patient Care", "Regulatory Compliance", "Data Analysis", "Quality Improvement", "Communication", "Leadership", "Process Automation", "AI/ML Tools", "Change Management", "Critical Thinking", "Project Management", "Stakeholder Mgmt", "Digital Health"],
        },
        "manufacturing": {
            "label": "Manufacturing",
            "company": "Pinnacle Industrial Corp",
            "functions": {
                "Production": {"roles": ["Production Manager", "Line Supervisor", "Quality Inspector", "Process Engineer", "Production Planner"], "comp_base": 75000},
                "Engineering": {"roles": ["Design Engineer", "Manufacturing Engineer", "Industrial Engineer", "Automation Engineer", "R&D Engineer"], "comp_base": 95000},
                "Supply Chain": {"roles": ["Supply Chain Manager", "Procurement Specialist", "Logistics Coordinator", "Inventory Analyst", "Demand Planner"], "comp_base": 80000},
                "Quality": {"roles": ["Quality Manager", "QA Engineer", "Six Sigma Lead", "Metrology Technician", "Compliance Specialist"], "comp_base": 85000},
                "Maintenance": {"roles": ["Maintenance Manager", "Reliability Engineer", "Maintenance Technician", "Predictive Maintenance Analyst", "Facilities Coordinator"], "comp_base": 70000},
                "EHS": {"roles": ["EHS Manager", "Safety Engineer", "Environmental Specialist", "Industrial Hygienist", "Sustainability Lead"], "comp_base": 80000},
                "IT / OT": {"roles": ["IT Manager", "OT Engineer", "SCADA Specialist", "MES Administrator", "Data Analyst"], "comp_base": 90000},
                "Finance": {"roles": ["Plant Controller", "Cost Accountant", "FP&A Analyst", "Payroll Specialist", "AP/AR Clerk"], "comp_base": 75000},
                "HR": {"roles": ["HR Manager", "Recruiter", "Training Coordinator", "Labor Relations", "Benefits Administrator"], "comp_base": 72000},
                "Sales": {"roles": ["Sales Manager", "Account Manager", "Technical Sales", "Customer Service Rep", "Pricing Analyst"], "comp_base": 85000},
            },
            "tasks_by_role": {},
            "skills": ["Lean Manufacturing", "Six Sigma", "Process Optimization", "Quality Systems", "ERP/MES", "Data Analysis", "Automation", "PLC Programming", "Supply Chain Mgmt", "Communication", "Leadership", "Project Management", "Safety Compliance", "CAD/CAM", "AI/ML Tools"],
        },
        "retail": {
            "label": "Retail & Consumer",
            "company": "Crest Commerce Group",
            "functions": {
                "Merchandising": {"roles": ["Buyer", "Merchandiser", "Category Manager", "Assortment Planner", "Pricing Analyst"], "comp_base": 85000},
                "Store Operations": {"roles": ["District Manager", "Store Manager", "Assistant Manager", "Department Lead", "Visual Merchandiser"], "comp_base": 65000},
                "E-Commerce": {"roles": ["E-Commerce Manager", "Digital Product Manager", "UX Designer", "Web Developer", "SEO Specialist"], "comp_base": 100000},
                "Supply Chain": {"roles": ["Supply Chain Director", "Distribution Manager", "Inventory Planner", "Logistics Analyst", "Warehouse Manager"], "comp_base": 80000},
                "Marketing": {"roles": ["CMO", "Brand Manager", "Social Media Manager", "CRM Manager", "Content Creator"], "comp_base": 90000},
                "Customer Experience": {"roles": ["CX Director", "Customer Service Manager", "Contact Center Lead", "Voice of Customer Analyst", "CX Designer"], "comp_base": 75000},
                "Finance": {"roles": ["Controller", "Financial Analyst", "Revenue Analyst", "Lease Accountant", "Treasury Analyst"], "comp_base": 85000},
                "HR": {"roles": ["CHRO", "HR Business Partner", "Talent Acquisition", "L&D Manager", "Compensation Analyst"], "comp_base": 80000},
                "Technology": {"roles": ["CTO", "Software Engineer", "Data Engineer", "POS Systems Admin", "IT Support Lead"], "comp_base": 110000},
                "Real Estate": {"roles": ["Real Estate Director", "Site Selector", "Lease Manager", "Construction Manager", "Facilities Coordinator"], "comp_base": 90000},
            },
            "tasks_by_role": {},
            "skills": ["Merchandising", "Demand Planning", "Customer Analytics", "E-Commerce", "Supply Chain", "POS Systems", "Data Analysis", "Communication", "Leadership", "Negotiation", "Digital Marketing", "AI/ML Tools", "Process Automation", "Visual Merchandising", "CRM"],
        },
        "professional_services": {
            "label": "Professional Services",
            "company": "Apex Advisory Partners",
            "functions": {
                "Consulting": {"roles": ["Managing Director", "Principal", "Senior Consultant", "Consultant", "Analyst"], "comp_base": 130000},
                "Advisory": {"roles": ["Partner", "Director", "Manager", "Senior Associate", "Associate"], "comp_base": 125000},
                "Delivery": {"roles": ["Engagement Manager", "Project Lead", "Delivery Lead", "Technical Architect", "Solutions Designer"], "comp_base": 120000},
                "Research": {"roles": ["Research Director", "Senior Researcher", "Industry Analyst", "Knowledge Manager", "Insights Lead"], "comp_base": 110000},
                "Business Development": {"roles": ["BD Director", "Proposal Manager", "Client Partner", "Account Director", "Pursuit Lead"], "comp_base": 115000},
                "People & Talent": {"roles": ["HR Director", "Recruiter", "L&D Lead", "Performance Manager", "Alumni Relations"], "comp_base": 95000},
                "Finance": {"roles": ["CFO", "Controller", "Billing Manager", "Revenue Analyst", "FP&A Analyst"], "comp_base": 100000},
                "Marketing": {"roles": ["CMO", "Content Director", "Events Manager", "Digital Marketing", "Brand Strategist"], "comp_base": 100000},
                "Technology": {"roles": ["CTO", "Solutions Architect", "Data Engineer", "Platform Lead", "Innovation Lead"], "comp_base": 130000},
                "Operations": {"roles": ["COO", "Office Manager", "Procurement Lead", "Facilities Manager", "Travel Coordinator"], "comp_base": 80000},
            },
            "tasks_by_role": {},
            "skills": ["Strategy", "Client Management", "Data Analysis", "Financial Modeling", "Project Management", "Communication", "Leadership", "Problem Solving", "Industry Expertise", "AI/ML Tools", "Presentation", "Research", "Change Management", "Stakeholder Mgmt", "Digital Transformation"],
        },
    }
    
    cfg = industry_configs.get(industry, industry_configs["technology"])
    
    # ═══ GENERATE EMPLOYEES ═══
    first_names = ["Sarah","James","Maria","David","Lisa","Michael","Emma","Robert","Jennifer","William","Amanda","Christopher","Nicole","Daniel","Rachel","Thomas","Catherine","Steven","Laura","Andrew","Olivia","Ethan","Sophia","Liam","Isabella","Noah","Mia","Lucas","Charlotte","Mason","Amelia","Logan","Harper","Alexander","Evelyn","Jacob","Abigail","Benjamin","Emily","Henry","Ella","Sebastian","Avery","Jack","Scarlett","Owen","Grace","Samuel","Chloe"]
    last_names = ["Chen","Rodriguez","Thompson","Kim","Patel","Brown","Wilson","Taylor","Garcia","Davis","Martinez","Lee","Johnson","White","Anderson","Harris","Clark","Wright","King","Scott","Adams","Baker","Mitchell","Rivera","Campbell","Phillips","Evans","Turner","Torres","Collins","Stewart","Morris","Murphy","Cook","Rogers","Morgan","Reed","Bell","Cooper","Bailey"]
    
    employees = []
    functions_list = list(cfg["functions"].keys())[:sc["functions"]]
    emp_id = 1
    managers = {}
    
    # Create org structure: C-suite → VPs → Directors → Managers → ICs
    levels = ["L1","L2","L3","L4","L5","L6","L7"][:sc["levels"]]
    
    # CEO
    ceo_id = f"EMP{str(emp_id).zfill(4)}"
    employees.append({"Employee ID": ceo_id, "Name": f"{random.choice(first_names)} {random.choice(last_names)}", "Job Title": "CEO", "Function": "Executive", "Sub-Function": "C-Suite", "Career Level": levels[-1], "Career Track": "Manager", "Manager ID": "", "Compensation": int(cfg["functions"][functions_list[0]]["comp_base"] * 3.5 * sc["comp_mult"]), "Tenure": random.randint(5, 15), "Location": random.choice(["New York", "San Francisco", "Chicago", "London"])})
    emp_id += 1
    
    # VPs per function
    for func in functions_list:
        vp_id = f"EMP{str(emp_id).zfill(4)}"
        vp_level = levels[-2] if len(levels) >= 2 else levels[-1]
        comp_base = cfg["functions"][func]["comp_base"]
        employees.append({"Employee ID": vp_id, "Name": f"{random.choice(first_names)} {random.choice(last_names)}", "Job Title": f"VP {func}", "Function": func, "Sub-Function": "Leadership", "Career Level": vp_level, "Career Track": "Manager", "Manager ID": ceo_id, "Compensation": int(comp_base * 2.0 * sc["comp_mult"]), "Tenure": random.randint(4, 12), "Location": random.choice(["New York", "San Francisco", "Chicago", "London", "Austin"])})
        managers[func] = vp_id
        emp_id += 1
    
    # Fill remaining employees across functions
    remaining = num_emp - len(employees)
    per_func = remaining // len(functions_list)
    
    for func in functions_list:
        roles = cfg["functions"][func]["roles"]
        comp_base = cfg["functions"][func]["comp_base"]
        mgr_id = managers.get(func, ceo_id)
        
        # Directors (if large enough)
        num_directors = max(1, per_func // 15) if sc["employees"] >= 200 else 0
        director_ids = []
        for _ in range(num_directors):
            d_id = f"EMP{str(emp_id).zfill(4)}"
            d_level = levels[-3] if len(levels) >= 3 else levels[-2]
            employees.append({"Employee ID": d_id, "Name": f"{random.choice(first_names)} {random.choice(last_names)}", "Job Title": f"Director, {func}", "Function": func, "Sub-Function": "Leadership", "Career Level": d_level, "Career Track": "Manager", "Manager ID": mgr_id, "Compensation": int(comp_base * 1.5 * sc["comp_mult"]), "Tenure": random.randint(3, 10), "Location": random.choice(["New York", "San Francisco", "Chicago", "London", "Austin", "Seattle"])})
            director_ids.append(d_id)
            emp_id += 1
        
        # ICs
        num_ics = per_func - num_directors
        for i in range(num_ics):
            if emp_id > num_emp:
                break
            role = roles[i % len(roles)]
            ic_level_idx = random.choices(range(min(3, len(levels))), weights=[30, 50, 20][:min(3, len(levels))])[0]
            report_to = random.choice(director_ids) if director_ids else mgr_id
            comp = int(comp_base * (0.7 + ic_level_idx * 0.25) * sc["comp_mult"] * random.uniform(0.85, 1.15))
            employees.append({"Employee ID": f"EMP{str(emp_id).zfill(4)}", "Name": f"{random.choice(first_names)} {random.choice(last_names)}", "Job Title": role, "Function": func, "Sub-Function": func, "Career Level": levels[ic_level_idx], "Career Track": random.choice(["IC", "IC", "IC", "Manager"]), "Manager ID": report_to, "Compensation": comp, "Tenure": random.randint(0, 8), "Location": random.choice(["New York", "San Francisco", "Chicago", "London", "Austin", "Seattle", "Denver", "Boston"])})
            emp_id += 1
    
    # ═══ GENERATE WORK DESIGN TASKS ═══
    task_types = ["Repetitive", "Variable"]
    logics = ["Deterministic", "Probabilistic", "Judgment-heavy"]
    interactions = ["Independent", "Interactive", "Collaborative"]
    impacts = ["High", "Moderate", "Low"]
    
    work_design_rows = []
    # Generate tasks for top 2 roles per function
    for func in functions_list[:4]:
        roles = cfg["functions"][func]["roles"][:2]
        for role in roles:
            # Check if we have predefined tasks
            if role in cfg.get("tasks_by_role", {}):
                for t_name, t_pct, t_type, t_logic, t_inter, t_impact, t_pskill, t_sskill in cfg["tasks_by_role"][role]:
                    work_design_rows.append({"Job Title": role, "Task Name": t_name, "Est Hours/Week": round(40 * t_pct / 100, 1), "Current Time Spent %": t_pct, "Task Type": t_type, "Logic": t_logic, "Interaction": t_inter, "AI Impact": t_impact, "Primary Skill": t_pskill, "Secondary Skill": t_sskill})
            else:
                # Generate generic tasks
                task_templates = [
                    (f"Data analysis & reporting for {func}", 20, "Repetitive", "Deterministic", "Independent", "High"),
                    (f"Stakeholder communication in {func}", 15, "Variable", "Judgment-heavy", "Collaborative", "Low"),
                    (f"Process documentation & compliance", 15, "Repetitive", "Deterministic", "Independent", "High"),
                    (f"Strategic planning & decision support", 20, "Variable", "Probabilistic", "Interactive", "Moderate"),
                    (f"Team coordination & project management", 15, "Variable", "Judgment-heavy", "Collaborative", "Low"),
                    (f"Quality review & validation", 15, "Repetitive", "Probabilistic", "Interactive", "Moderate"),
                ]
                for t_name, t_pct, t_type, t_logic, t_inter, t_impact in task_templates:
                    skills = cfg["skills"]
                    work_design_rows.append({"Job Title": role, "Task Name": t_name, "Est Hours/Week": round(40 * t_pct / 100, 1), "Current Time Spent %": t_pct, "Task Type": t_type, "Logic": t_logic, "Interaction": t_inter, "AI Impact": t_impact, "Primary Skill": random.choice(skills[:5]), "Secondary Skill": random.choice(skills[5:10])})
    
    # ═══ GENERATE SKILLS DATA ═══
    skills_rows = []
    for emp in employees[:min(num_emp, 200)]:  # Cap at 200 for performance
        for skill in cfg["skills"]:
            if random.random() > 0.2:
                base = 3 if skill in cfg["skills"][:3] and emp["Career Level"] in levels[-3:] else 2
                prof = max(1, min(4, base + random.choices([-1, 0, 0, 1], weights=[10, 35, 35, 20])[0]))
                skills_rows.append({"Employee ID": emp["Employee ID"], "Name": emp["Name"], "Skill": skill, "Proficiency": prof})
    
    # ═══ BUILD DATAFRAMES ═══
    wf_df = pd.DataFrame(employees)
    wd_df = pd.DataFrame(work_design_rows) if work_design_rows else pd.DataFrame()
    sk_df = pd.DataFrame(skills_rows) if skills_rows else pd.DataFrame()
    
    # Operating model
    op_rows = []
    for func in functions_list[:6]:
        for cap in ["Strategy & Planning", "Execution & Delivery", "Analytics & Reporting", "Governance & Compliance", "Innovation & Transformation"]:
            op_rows.append({"Function": func, "Capability": f"{func} {cap}", "Maturity": random.randint(1, 4), "Service Model": random.choice(["Shared", "Embedded", "CoE"])})
    op_df = pd.DataFrame(op_rows)
    
    # Change management
    ch_rows = []
    for emp in employees[:min(num_emp, 100)]:
        ch_rows.append({"Employee ID": emp["Employee ID"], "Name": emp["Name"], "Function": emp["Function"], "Change Readiness": round(random.uniform(1.5, 4.5), 1), "Impact Level": random.choice(["High", "Moderate", "Low"]), "Support Needed": random.choice(["Intensive", "Moderate", "Light"])})
    ch_df = pd.DataFrame(ch_rows)
    
    model_id = f"Tutorial_{industry}_{size}"
    
    return {
        "model_id": model_id,
        "company": cfg["company"],
        "industry": cfg["label"],
        "size_label": sc["label"],
        "datasets": {
            "workforce": wf_df,
            "work_design": wd_df,
            "skills": sk_df,
            "org_design": wf_df[["Employee ID", "Name", "Job Title", "Function", "Career Level", "Manager ID"]].copy(),
            "operating_model": op_df,
            "change_management": ch_df,
        }
    }


def _seed_tutorial_store(industry="technology", size_tier="mid"):
    """Create full-scale org dataset. size_tier: small, mid, large."""
    import pandas as pd
    import numpy as np
    import random
    
    model_id = f"Tutorial_{size_tier.title()}_{industry.title().replace(' ', '_')}"
    
    # Always regenerate — don't cache, to ensure schema changes are picked up
    random.seed(hash(f"{industry}_{size_tier}") % 2**31)

    # ═══════════════════════════════════════════════════
    # 24 COMPANIES — 8 industries × 3 sizes
    # ═══════════════════════════════════════════════════
    
    COMPANY_DB = {
        "technology": {
            "small":  {"name": "Spark Labs", "employees": 150, "revenue": "$25M", "hq": "Austin"},
            "mid":    {"name": "Nexus Technology Corp", "employees": 2500, "revenue": "$450M", "hq": "San Francisco"},
            "large":  {"name": "Titan Digital Systems", "employees": 18000, "revenue": "$4.2B", "hq": "Seattle"},
        },
        "financial_services": {
            "small":  {"name": "Pinnacle Wealth Advisors", "employees": 200, "revenue": "$80M AUM", "hq": "Charlotte"},
            "mid":    {"name": "Global Financial Partners", "employees": 4200, "revenue": "$1.8B", "hq": "New York"},
            "large":  {"name": "Meridian Capital Group", "employees": 22000, "revenue": "$12B", "hq": "New York"},
        },
        "healthcare": {
            "small":  {"name": "Valley Medical Center", "employees": 350, "revenue": "$120M", "hq": "Portland"},
            "mid":    {"name": "Meridian Health System", "employees": 5500, "revenue": "$2.1B", "hq": "Chicago"},
            "large":  {"name": "National Health Partners", "employees": 45000, "revenue": "$18B", "hq": "Nashville"},
        },
        "manufacturing": {
            "small":  {"name": "Precision Components Inc", "employees": 180, "revenue": "$35M", "hq": "Cincinnati"},
            "mid":    {"name": "Atlas Manufacturing Group", "employees": 3800, "revenue": "$950M", "hq": "Detroit"},
            "large":  {"name": "Continental Industrial Corp", "employees": 28000, "revenue": "$8.5B", "hq": "Milwaukee"},
        },
        "retail": {
            "small":  {"name": "Urban Threads Boutique", "employees": 120, "revenue": "$18M", "hq": "Denver"},
            "mid":    {"name": "Horizon Retail Group", "employees": 6000, "revenue": "$1.4B", "hq": "Minneapolis"},
            "large":  {"name": "American Marketplace Inc", "employees": 52000, "revenue": "$22B", "hq": "Dallas"},
        },
        "legal": {
            "small":  {"name": "Barrett & Associates", "employees": 45, "revenue": "$12M", "hq": "Boston"},
            "mid":    {"name": "Sterling Legal Group", "employees": 800, "revenue": "$280M", "hq": "Washington DC"},
            "large":  {"name": "Global Law Alliance", "employees": 5500, "revenue": "$3.2B", "hq": "London"},
        },
        "energy": {
            "small":  {"name": "SunRidge Renewables", "employees": 160, "revenue": "$45M", "hq": "Denver"},
            "mid":    {"name": "Apex Energy Solutions", "employees": 3200, "revenue": "$2.8B", "hq": "Houston"},
            "large":  {"name": "Pacific Energy Holdings", "employees": 35000, "revenue": "$28B", "hq": "Houston"},
        },
        "education": {
            "small":  {"name": "Westbrook College", "employees": 120, "revenue": "$28M", "hq": "Portland"},
            "mid":    {"name": "Pacific University System", "employees": 2800, "revenue": "$650M", "hq": "Seattle"},
            "large":  {"name": "National University Consortium", "employees": 15000, "revenue": "$4.5B", "hq": "Philadelphia"},
        },
    }

    company = COMPANY_DB.get(industry, COMPANY_DB["technology"]).get(size_tier, COMPANY_DB["technology"]["mid"])
    target_size = company["employees"]
    
    # Generate the actual target size for each tier
    # Small: full generation (45-350 employees)  
    # Mid: full generation (800-6000 employees)
    # Large: representative sample capped at 8000 for performance (still 4-8x larger than mid)
    gen_size = min(target_size, 8000)
    
    # ═══════════════════════════════════════════════════
    # ORGANIZATIONAL BLUEPRINT — 9 functions per industry
    # ═══════════════════════════════════════════════════
    
    FUNC_BLUEPRINTS = {
        "technology": [
            ("Engineering", 0.30, [("Junior Software Engineer","L2","IC",85000,105000,0.30), ("Software Engineer","L3","IC",110000,140000,0.25), ("Senior Software Engineer","L4","IC",140000,180000,0.20), ("Staff Engineer","L5","IC",180000,220000,0.08), ("Engineering Manager","L4","Manager",150000,190000,0.10), ("Sr Engineering Manager","L5","Manager",190000,230000,0.05), ("VP Engineering","L6","Manager",230000,300000,0.02)]),
            ("Product", 0.12, [("Associate PM","L2","IC",90000,110000,0.25), ("Product Manager","L3","IC",120000,150000,0.30), ("Senior PM","L4","IC",150000,190000,0.20), ("UX Designer","L3","IC",95000,125000,0.15), ("Director of Product","L5","Manager",190000,240000,0.08), ("VP Product","L6","Manager",240000,300000,0.02)]),
            ("Data & Analytics", 0.15, [("Data Analyst","L3","IC",80000,100000,0.30), ("Senior Data Analyst","L4","IC",100000,130000,0.20), ("Data Engineer","L3","IC",110000,140000,0.20), ("Data Scientist","L4","IC",130000,170000,0.12), ("ML Engineer","L4","IC",140000,180000,0.08), ("Analytics Manager","L4","Manager",130000,160000,0.07), ("VP Data","L6","Manager",220000,280000,0.03)]),
            ("IT Operations", 0.10, [("IT Support","L2","IC",55000,72000,0.30), ("Sys Admin","L3","IC",75000,95000,0.25), ("DevOps Engineer","L3","IC",110000,140000,0.20), ("SRE","L4","IC",130000,165000,0.12), ("IT Manager","L4","Manager",120000,150000,0.08), ("Director IT","L5","Manager",160000,200000,0.05)]),
            ("Security", 0.06, [("Security Analyst","L3","IC",90000,115000,0.35), ("Security Engineer","L4","IC",120000,155000,0.30), ("Pen Tester","L3","IC",100000,130000,0.15), ("Security Manager","L4","Manager",140000,175000,0.12), ("CISO","L6","Manager",210000,280000,0.08)]),
            ("Sales & Marketing", 0.12, [("SDR","L2","IC",55000,70000,0.25), ("Account Executive","L3","IC",80000,120000,0.25), ("Marketing Specialist","L3","IC",70000,90000,0.20), ("Content Strategist","L3","IC",75000,95000,0.10), ("Sales Manager","L4","Manager",120000,160000,0.10), ("VP Marketing","L6","Manager",200000,260000,0.05), ("VP Sales","L6","Manager",210000,280000,0.05)]),
            ("Finance", 0.06, [("Financial Analyst","L3","IC",75000,95000,0.30), ("Accountant","L3","IC",65000,85000,0.25), ("FP&A Manager","L4","Manager",110000,140000,0.15), ("Controller","L5","Manager",150000,190000,0.10), ("AP/AR Specialist","L2","IC",48000,62000,0.15), ("CFO","L6","Manager",240000,320000,0.05)]),
            ("HR & People", 0.05, [("HR Coordinator","L2","IC",50000,65000,0.25), ("HRBP","L3","IC",80000,100000,0.25), ("Recruiter","L3","IC",70000,90000,0.20), ("L&D Specialist","L3","IC",72000,92000,0.10), ("HR Manager","L4","Manager",100000,130000,0.10), ("VP People","L6","Manager",200000,260000,0.05), ("Comp Analyst","L3","IC",72000,92000,0.05)]),
            ("Legal", 0.04, [("Paralegal","L2","IC",55000,72000,0.30), ("Corporate Counsel","L4","IC",140000,180000,0.25), ("Compliance Analyst","L3","IC",80000,100000,0.20), ("Privacy Officer","L4","IC",120000,155000,0.10), ("General Counsel","L6","Manager",230000,300000,0.05), ("Contract Manager","L3","IC",85000,110000,0.10)]),
        ],
    }
    
    # Industry-specific function & role overrides
    RENAMES = {
        "financial_services": {"Engineering":"Technology","Product":"Investment Banking","Data & Analytics":"Risk Management","IT Operations":"Operations","Security":"Compliance","Sales & Marketing":"Wealth Management","Finance":"Treasury","HR & People":"HR","Legal":"Legal & Regulatory",
            "Junior Software Engineer":"Junior Analyst","Software Engineer":"Quantitative Analyst","Senior Software Engineer":"Senior Quant","Staff Engineer":"VP Quant Research","Engineering Manager":"Tech Lead","Associate PM":"Investment Analyst","Product Manager":"Associate","Senior PM":"VP Investments","UX Designer":"Trading Analyst","SDR":"Client Associate","Account Executive":"Relationship Manager","Data Analyst":"Risk Analyst","Accountant":"Fund Accountant","Compliance Analyst":"Regulatory Analyst","IT Support":"Operations Associate","Recruiter":"TA Specialist","Paralegal":"Legal Analyst"},
        "healthcare": {"Engineering":"Health Informatics","Product":"Clinical Operations","Data & Analytics":"Revenue Cycle","IT Operations":"IT Services","Security":"Quality & Safety","Sales & Marketing":"Patient Services","Legal":"Compliance",
            "Junior Software Engineer":"Clinical Systems Analyst","Software Engineer":"EHR Developer","Data Analyst":"Clinical Data Analyst","Product Manager":"Care Coordinator","SDR":"Patient Access Rep","Account Executive":"Provider Relations","UX Designer":"Medical Coder","IT Support":"Help Desk","Security Analyst":"Quality Auditor","Compliance Analyst":"HIPAA Analyst","Recruiter":"Nurse Recruiter","Accountant":"Revenue Cycle Specialist"},
        "manufacturing": {"Product":"Production","Data & Analytics":"Quality Control","IT Operations":"Maintenance","Security":"EHS","Sales & Marketing":"Supply Chain","Legal":"Regulatory",
            "Software Engineer":"Process Engineer","Data Analyst":"Quality Analyst","Product Manager":"Production Planner","SDR":"Procurement Specialist","Account Executive":"Supply Chain Analyst","UX Designer":"Industrial Engineer","IT Support":"Maintenance Technician","Security Analyst":"EHS Coordinator","DevOps Engineer":"Automation Engineer"},
        "retail": {"Engineering":"E-Commerce Tech","Product":"Merchandising","Data & Analytics":"Analytics","IT Operations":"Store Operations","Security":"Loss Prevention","Sales & Marketing":"Marketing",
            "Software Engineer":"E-Commerce Developer","Data Analyst":"Demand Planner","Product Manager":"Category Manager","SDR":"Store Associate","Account Executive":"District Manager","UX Designer":"Visual Merchandiser","IT Support":"Store Tech","Security Analyst":"LP Analyst"},
        "legal": {"Engineering":"Legal Technology","Product":"Litigation","Data & Analytics":"Knowledge Management","IT Operations":"Legal Ops","Security":"Risk & Compliance","Sales & Marketing":"Business Development",
            "Software Engineer":"Legal Tech Developer","Data Analyst":"eDiscovery Specialist","Product Manager":"Practice Group Lead","SDR":"BD Coordinator","Account Executive":"Client Partner","UX Designer":"Knowledge Manager","Security Analyst":"Regulatory Analyst","Accountant":"Billing Coordinator"},
        "energy": {"Product":"Operations","Data & Analytics":"Asset Management","IT Operations":"SCADA/OT","Security":"HSE","Sales & Marketing":"Trading & Commercial",
            "Software Engineer":"Controls Engineer","Data Analyst":"Asset Integrity Analyst","Product Manager":"Operations Superintendent","SDR":"Energy Trader","Account Executive":"Commercial Manager","UX Designer":"Renewables Analyst","DevOps Engineer":"SCADA Engineer","Security Analyst":"HSE Coordinator"},
        "education": {"Engineering":"IT Services","Product":"Academic Affairs","Data & Analytics":"Institutional Research","IT Operations":"IT Operations","Security":"Student Services","Sales & Marketing":"Enrollment",
            "Software Engineer":"Systems Developer","Data Analyst":"IR Analyst","Product Manager":"Curriculum Coordinator","SDR":"Admissions Counselor","Account Executive":"Enrollment Manager","UX Designer":"Instructional Designer","Security Analyst":"Student Advisor","Recruiter":"Faculty Recruiter"},
    }
    
    INDUSTRY_SKILLS = {
        "technology": ["Data Analysis","Python/SQL","Cloud Platforms","AI/ML Tools","Process Automation","Communication","Leadership","Stakeholder Mgmt","Critical Thinking","Digital Fluency","AI Literacy","Change Agility","Cybersecurity","DevOps","UX Design"],
        "financial_services": ["Data Analysis","Financial Modeling","Risk Assessment","Regulatory Knowledge","Process Automation","Communication","Leadership","Stakeholder Mgmt","Critical Thinking","AI/ML Tools","Python/SQL","Compliance","Negotiation","Digital Fluency","AI Literacy"],
        "healthcare": ["Clinical Data Analysis","Medical Coding","EHR Systems","Regulatory Knowledge","Process Automation","Communication","Leadership","Stakeholder Mgmt","Critical Thinking","AI/ML Tools","Patient Experience","Quality Improvement","HIPAA Compliance","Digital Fluency","AI Literacy"],
        "manufacturing": ["Data Analysis","Process Engineering","Quality Control","Lean/Six Sigma","Supply Chain Mgmt","Communication","Leadership","Stakeholder Mgmt","Critical Thinking","AI/ML Tools","ERP Systems","Predictive Maintenance","Safety Compliance","Digital Fluency","AI Literacy"],
        "retail": ["Data Analysis","Demand Forecasting","Inventory Mgmt","Customer Experience","Process Automation","Communication","Leadership","Stakeholder Mgmt","Critical Thinking","AI/ML Tools","E-Commerce","Digital Marketing","Negotiation","Digital Fluency","AI Literacy"],
        "legal": ["Legal Research","Contract Analysis","eDiscovery","Regulatory Knowledge","Process Automation","Communication","Leadership","Stakeholder Mgmt","Critical Thinking","AI/ML Tools","IP Management","Negotiation","Knowledge Mgmt","Digital Fluency","AI Literacy"],
        "energy": ["Data Analysis","Process Engineering","Asset Mgmt","HSE Knowledge","Process Automation","Communication","Leadership","Stakeholder Mgmt","Critical Thinking","AI/ML Tools","SCADA Systems","Predictive Analytics","Regulatory Knowledge","Digital Fluency","AI Literacy"],
        "education": ["Data Analysis","Student Info Systems","Enrollment Mgmt","Financial Aid","Process Automation","Communication","Leadership","Stakeholder Mgmt","Critical Thinking","AI/ML Tools","Academic Advising","Grant Mgmt","Curriculum Design","Digital Fluency","AI Literacy"],
    }

    FIRST_NAMES = ["Sarah","James","Maria","David","Lisa","Michael","Emma","Robert","Jennifer","William","Amanda","Christopher","Nicole","Daniel","Rachel","Thomas","Catherine","Steven","Laura","Andrew","Jessica","Kevin","Michelle","Brandon","Stephanie","Tyler","Ashley","Nathan","Brittany","Ryan","Megan","Justin","Emily","Aaron","Heather","Patrick","Rebecca","Sean","Samantha","Jacob","Amber","Eric","Kayla","Brian","Lauren","Gregory","Kelly","Mark","Tiffany","Jason","Olivia","Ethan","Hannah","Dylan","Sophia","Logan","Ava","Mason","Isabella","Lucas","Mia","Noah","Abigail","Liam","Charlotte","Aiden","Ella"]
    LAST_NAMES = ["Chen","Rodriguez","Thompson","Kim","Patel","Brown","Wilson","Taylor","Garcia","Davis","Martinez","Lee","Johnson","White","Anderson","Harris","Clark","Wright","King","Scott","Moore","Jackson","Martin","Lewis","Allen","Walker","Young","Hall","Hernandez","Lopez","Hill","Green","Adams","Baker","Gonzalez","Nelson","Mitchell","Perez","Roberts","Turner","Phillips","Campbell","Parker","Evans","Edwards","Collins","Stewart","Sanchez","Morris"]
    LOCATIONS = ["New York","Chicago","San Francisco","Boston","Austin","London","Toronto","Seattle","Denver","Atlanta","Houston","Miami","Los Angeles","Dallas","Phoenix"]

    renames = RENAMES.get(industry, {})
    skills = INDUSTRY_SKILLS.get(industry, INDUSTRY_SKILLS["technology"])
    blueprints = FUNC_BLUEPRINTS["technology"]  # Use tech as base

    emp_id = 1
    employees = []
    mgr_map = {}

    for base_func, pct, roles in blueprints:
        real_func = renames.get(base_func, base_func)
        func_count = max(2, round(gen_size * pct))
        func_mgrs = []
        
        for title, level, track, comp_lo, comp_hi, role_pct in roles:
            real_title = renames.get(title, title)
            role_count = max(1, round(func_count * role_pct))
            
            for _ in range(role_count):
                name = f"{random.choice(FIRST_NAMES)} {random.choice(LAST_NAMES)}"
                comp = random.randint(comp_lo, comp_hi)
                eid = f"EMP{str(emp_id).zfill(5)}"
                
                employees.append({
                    "Employee ID": eid, "Name": name, "Job Title": real_title,
                    "Function": real_func, "Sub-Function": real_func,
                    "Career Level": level, "Career Track": track,
                    "Manager ID": "", "Compensation": comp,
                    "Tenure": random.randint(0, 20),
                    "Location": random.choice(LOCATIONS),
                })
                if track == "Manager": func_mgrs.append(eid)
                emp_id += 1
        
        mgr_map[real_func] = func_mgrs

    # Assign managers
    for emp in employees:
        if emp["Career Level"] in ["L5","L6"]: continue
        mgrs = mgr_map.get(emp["Function"], [])
        if mgrs: emp["Manager ID"] = random.choice(mgrs)

    # TASK LIBRARY — comprehensive per-industry
    tasks = []
    task_templates = [
        ("Data Analyst|Risk Analyst|Clinical Data|Quality Analyst|Demand Planner|eDiscovery|Asset Integrity|IR Analyst", [
            ("Data extraction & pipeline maintenance", 25, "Repetitive", "Deterministic", "Independent", "High"),
            ("Report generation & formatting", 20, "Repetitive", "Deterministic", "Independent", "High"),
            ("Ad-hoc stakeholder analysis", 15, "Variable", "Probabilistic", "Interactive", "Moderate"),
            ("Dashboard creation & maintenance", 10, "Repetitive", "Deterministic", "Independent", "High"),
            ("Executive presentations", 10, "Variable", "Judgment-heavy", "Collaborative", "Low"),
            ("Data quality monitoring", 8, "Repetitive", "Deterministic", "Independent", "High"),
            ("Cross-team consulting", 7, "Variable", "Judgment-heavy", "Collaborative", "Low"),
            ("Tool evaluation", 5, "Variable", "Probabilistic", "Interactive", "Moderate"),
        ]),
        ("Software Engineer|Developer|Process Engineer|Controls Engineer|Systems Developer|EHR Developer|Legal Tech", [
            ("Feature development & coding", 30, "Variable", "Judgment-heavy", "Collaborative", "Low"),
            ("Code review & PR management", 15, "Variable", "Judgment-heavy", "Collaborative", "Moderate"),
            ("Writing automated tests", 15, "Repetitive", "Deterministic", "Independent", "High"),
            ("Bug triage & debugging", 15, "Variable", "Probabilistic", "Independent", "Moderate"),
            ("Documentation writing", 10, "Repetitive", "Deterministic", "Independent", "High"),
            ("Sprint planning & estimation", 10, "Variable", "Judgment-heavy", "Collaborative", "Low"),
            ("Incident response", 5, "Variable", "Probabilistic", "Interactive", "Moderate"),
        ]),
        ("Financial Analyst|Accountant|Fund Accountant|Revenue Cycle|Billing|AP/AR", [
            ("Monthly close & reconciliation", 25, "Repetitive", "Deterministic", "Independent", "High"),
            ("Variance analysis & reporting", 20, "Variable", "Probabilistic", "Interactive", "Moderate"),
            ("Invoice/payment processing", 15, "Repetitive", "Deterministic", "Independent", "High"),
            ("Budget forecasting", 15, "Variable", "Probabilistic", "Interactive", "Moderate"),
            ("Audit preparation", 10, "Variable", "Judgment-heavy", "Collaborative", "Low"),
            ("Expense report review", 10, "Repetitive", "Deterministic", "Independent", "High"),
            ("Stakeholder presentations", 5, "Variable", "Judgment-heavy", "Collaborative", "Low"),
        ]),
        ("Recruiter|HR Coordinator|TA Specialist|Nurse Recruiter|Faculty Recruiter", [
            ("Resume screening & shortlisting", 25, "Repetitive", "Deterministic", "Independent", "High"),
            ("Interview scheduling", 20, "Repetitive", "Deterministic", "Interactive", "High"),
            ("Candidate sourcing & outreach", 20, "Variable", "Judgment-heavy", "Interactive", "Moderate"),
            ("Offer preparation", 15, "Variable", "Judgment-heavy", "Collaborative", "Low"),
            ("Onboarding management", 10, "Repetitive", "Deterministic", "Interactive", "High"),
            ("Recruitment metrics", 10, "Repetitive", "Deterministic", "Independent", "High"),
        ]),
        ("SDR|Account Executive|Relationship Manager|Client|Store Associate|BD Coordinator|Admissions|Patient Access", [
            ("Prospecting & lead qualification", 25, "Variable", "Probabilistic", "Interactive", "Moderate"),
            ("CRM data entry & pipeline mgmt", 20, "Repetitive", "Deterministic", "Independent", "High"),
            ("Demo & presentation delivery", 20, "Variable", "Judgment-heavy", "Collaborative", "Low"),
            ("Proposal & contract prep", 15, "Repetitive", "Deterministic", "Independent", "High"),
            ("Competitive research", 10, "Variable", "Probabilistic", "Independent", "High"),
            ("Pipeline forecasting", 10, "Repetitive", "Deterministic", "Interactive", "High"),
        ]),
        ("Compliance|Regulatory|HIPAA|Quality Auditor|EHS Coordinator", [
            ("Regulatory monitoring & updates", 20, "Repetitive", "Deterministic", "Independent", "High"),
            ("Audit preparation & documentation", 20, "Repetitive", "Deterministic", "Independent", "High"),
            ("Policy review & updates", 15, "Variable", "Judgment-heavy", "Collaborative", "Low"),
            ("Investigation & case management", 15, "Variable", "Probabilistic", "Interactive", "Moderate"),
            ("Training & awareness programs", 15, "Variable", "Judgment-heavy", "Collaborative", "Low"),
            ("Reporting to regulators", 10, "Repetitive", "Deterministic", "Independent", "High"),
            ("Risk assessment", 5, "Variable", "Judgment-heavy", "Interactive", "Moderate"),
        ]),
        ("IT Support|Help Desk|Store Tech|Maintenance Tech|Sys Admin", [
            ("Ticket resolution & troubleshooting", 30, "Repetitive", "Deterministic", "Interactive", "High"),
            ("System monitoring & alerting", 20, "Repetitive", "Deterministic", "Independent", "High"),
            ("Software installation & updates", 15, "Repetitive", "Deterministic", "Independent", "High"),
            ("User onboarding & access mgmt", 15, "Repetitive", "Deterministic", "Interactive", "High"),
            ("Documentation & knowledge base", 10, "Repetitive", "Deterministic", "Independent", "High"),
            ("Escalation & vendor coordination", 10, "Variable", "Judgment-heavy", "Collaborative", "Low"),
        ]),
    ]

    import re
    matched_roles = set()
    for pattern, task_list in task_templates:
        for emp in employees:
            if re.search(pattern, emp["Job Title"], re.IGNORECASE):
                if emp["Job Title"] not in matched_roles:
                    matched_roles.add(emp["Job Title"])
                    for t_name, t_pct, t_type, t_logic, t_inter, t_impact in task_list:
                        tasks.append({
                            "Job Title": emp["Job Title"], "Task Name": t_name,
                            "Est Hours/Week": round(40 * t_pct / 100, 1),
                            "Current Time Spent %": t_pct,
                            "Task Type": t_type, "Logic": t_logic,
                            "Interaction": t_inter, "AI Impact": t_impact,
                            "Primary Skill": skills[0], "Secondary Skill": skills[4],
                        })

    # Generate generic tasks for unmatched roles — ensures larger orgs have broader WDL coverage
    all_titles = set(e["Job Title"] for e in employees)
    unmatched = all_titles - matched_roles
    generic_tasks = [
        ("Administrative coordination & scheduling", 15, "Repetitive", "Deterministic", "Interactive", "High"),
        ("Email & communication management", 15, "Repetitive", "Deterministic", "Independent", "High"),
        ("Data entry & record keeping", 15, "Repetitive", "Deterministic", "Independent", "High"),
        ("Reporting & presentation prep", 15, "Variable", "Probabilistic", "Interactive", "Moderate"),
        ("Cross-functional meetings & collaboration", 15, "Variable", "Judgment-heavy", "Collaborative", "Low"),
        ("Strategic planning & analysis", 15, "Variable", "Judgment-heavy", "Collaborative", "Low"),
        ("Process improvement & documentation", 10, "Variable", "Probabilistic", "Interactive", "Moderate"),
    ]
    # For larger orgs, cover more unmatched roles
    unmatched_limit = min(len(unmatched), 10 if gen_size < 500 else 25 if gen_size < 3000 else 50)
    for title in list(unmatched)[:unmatched_limit]:
        for t_name, t_pct, t_type, t_logic, t_inter, t_impact in generic_tasks:
            tasks.append({
                "Job Title": title, "Task Name": t_name,
                "Est Hours/Week": round(40 * t_pct / 100, 1),
                "Current Time Spent %": t_pct,
                "Task Type": t_type, "Logic": t_logic,
                "Interaction": t_inter, "AI Impact": t_impact,
                "Primary Skill": random.choice(skills[:5]), "Secondary Skill": random.choice(skills[5:10]),
            })

    # Generate skills — scale cap with company size
    # Small: all employees, Mid: up to 1000, Large: up to 2000
    skill_cap = min(len(employees), max(500, gen_size // 4))
    skill_emps = employees[:skill_cap]
    skills_records = []
    for emp in skill_emps:
        for skill in skills:
            if random.random() > 0.12:
                base_prof = 2
                if emp["Career Level"] in ["L4","L5","L6"]: base_prof = 3
                prof = max(1, min(4, base_prof + random.choices([-1,0,0,1], weights=[10,30,40,20])[0]))
                skills_records.append({"Employee ID": emp["Employee ID"], "Name": emp["Name"], "Skill": skill, "Proficiency": prof})

    # Operating model — more capabilities for larger orgs
    op_data = []
    funcs_used = list(set(e["Function"] for e in employees))
    op_caps_per_func = 5 if gen_size < 500 else 7 if gen_size < 3000 else 9
    base_caps = ["Strategy & Planning","Service Delivery","Analytics & Reporting","Governance & Risk","Innovation & Transformation","Talent Development","Process Excellence","Technology Enablement","Customer Experience"]
    for func in funcs_used[:min(len(funcs_used), 9)]:
        for cap in base_caps[:op_caps_per_func]:
            op_data.append({"Function": func, "Capability": f"{func} {cap}", "Maturity": random.randint(1,4), "Service Model": random.choice(["Shared","Embedded","CoE","Platform"])})

    # Change management — scale with org size
    change_cap = min(len(employees), max(100, gen_size // 3))
    change_emps = random.sample(employees, change_cap)
    change_data = [{"Employee ID": e["Employee ID"], "Name": e["Name"], "Function": e["Function"],
        "Change Readiness": round(random.uniform(1.5,4.5),1), "Impact Level": random.choice(["High","Moderate","Low"]),
        "Support Needed": random.choice(["Intensive","Moderate","Light"])} for e in change_emps]

    wf_df = pd.DataFrame(employees)
    wd_df = pd.DataFrame(tasks)
    sk_df = pd.DataFrame(skills_records)
    op_df = pd.DataFrame(op_data)
    ch_df = pd.DataFrame(change_data)

    # Add Model ID to all DataFrames (required by schemas)
    for df in [wf_df, wd_df, sk_df, op_df, ch_df]:
        if not df.empty:
            df["Model ID"] = model_id

    # Standardize all DataFrames through the schema pipeline so column names
    # match what API endpoints expect (e.g. "Name" → "Employee Name", 
    # "Function" → "Function ID", "Compensation" → "Base Pay")
    from app.store import standardize_to_schema
    wf_std = standardize_to_schema(wf_df, "workforce") if not wf_df.empty else pd.DataFrame()
    wd_std = standardize_to_schema(wd_df, "work_design") if not wd_df.empty else pd.DataFrame()
    
    # org_design uses workforce schema columns
    org_src = wf_df.copy()
    org_src["Model ID"] = model_id
    org_std = standardize_to_schema(org_src, "workforce") if not org_src.empty else pd.DataFrame()

    store.datasets[model_id] = {
        **empty_bundle(),
        "workforce": wf_std,
        "work_design": wd_std,
        "skills": sk_df,
        "org_design": org_std,
        "operating_model": op_df,
        "change_management": ch_df,
    }
    store.last_loaded_model_id = model_id
    
    print(f"Seeded: {model_id} — {company['name']} ({company['employees']:,} employees) → {len(employees)} generated, {len(tasks)} tasks, {len(skills_records)} skill records")
    return model_id

# Tutorial auto-seeds on /api/tutorial/seed call from frontend


@app.get("/api/tutorial/seed")
async def seed_tutorial(industry: str = "technology", size: str = "small"):
    """Seed tutorial sandbox with configurable industry and size."""
    try:
        size_map = {"small": "small", "mid": "mid", "medium": "mid", "large": "large", "mega": "large"}
        normalized_size = size_map.get(size, "mid")
        
        model_id = _seed_tutorial_store(industry, normalized_size)
        ds = store.datasets.get(model_id, {})
        emp_count = len(ds.get("workforce", [])) if isinstance(ds.get("workforce"), pd.DataFrame) else 0
        task_count = len(ds.get("work_design", [])) if isinstance(ds.get("work_design"), pd.DataFrame) and not ds["work_design"].empty else 0
        
        return _safe({"status": "ok", "model": model_id, "employees": emp_count, "tasks": task_count, "industry": industry, "size": normalized_size})
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(500, detail=str(e))

@app.get("/api/tutorial/industries")
async def list_industries():
    """List available industry × size combinations."""
    return {
        "industries": [
            {"id": "financial_services", "label": "Financial Services", "icon": "🏦", "example": "JP Morgan, Goldman Sachs"},
            {"id": "technology", "label": "Technology", "icon": "💻", "example": "Google, Salesforce"},
            {"id": "healthcare", "label": "Healthcare", "icon": "🏥", "example": "UnitedHealth, Kaiser"},
            {"id": "manufacturing", "label": "Manufacturing", "icon": "🏭", "example": "Toyota, Boeing"},
            {"id": "retail", "label": "Retail & Consumer", "icon": "🛍️", "example": "Walmart, Nike"},
            {"id": "professional_services", "label": "Professional Services", "icon": "💼", "example": "McKinsey, Deloitte"},
        ],
        "sizes": [
            {"id": "small", "label": "Small", "employees": "~50", "icon": "🏠"},
            {"id": "medium", "label": "Medium", "employees": "~200", "icon": "🏢"},
            {"id": "large", "label": "Large", "employees": "~1,000", "icon": "🏙️"},
            {"id": "mega", "label": "Mega", "employees": "~5,000", "icon": "🌍"},
        ],
    }


# ═══════════════════════════════════════════════════
# EXPORT — Real .docx report + Excel templates
# ═══════════════════════════════════════════════════

@app.get("/api/export/docx/{model_id}")
async def export_docx(model_id: str):
    """Generate a real .docx transformation report."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io, datetime
    
    # Gather all data
    summary = await get_export_summary(model_id)
    
    doc = Document()
    
    # Title
    title = doc.add_heading("AI Transformation Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(59, 130, 246)
    
    doc.add_paragraph(f"Generated: {datetime.datetime.now().strftime('%B %d, %Y')}")
    doc.add_paragraph(f"Model: {model_id}")
    doc.add_paragraph("")
    
    # Executive Summary
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(f"This report covers the AI transformation analysis for {model_id} with {summary.get('total_employees', 0)} employees assessed.")
    doc.add_paragraph(f"Organization AI Readiness: {summary.get('org_readiness', '—')}/5")
    doc.add_paragraph(f"Critical skill gaps identified: {summary.get('critical_gaps', 0)}")
    doc.add_paragraph(f"High-risk change segment: {summary.get('high_risk_pct', 0)}% of workforce")
    
    # Workforce Discovery
    doc.add_heading("2. Workforce Discovery", level=1)
    doc.add_paragraph(f"Total employees: {summary.get('total_employees', 0)}")
    doc.add_paragraph(f"Skills data coverage: {summary.get('skills_coverage', 0)}%")
    bands = summary.get("readiness_bands", {})
    if bands:
        doc.add_paragraph(f"Readiness bands: Ready Now ({bands.get('ready_now', 0)}), Coachable ({bands.get('coachable', 0)}), At Risk ({bands.get('at_risk', 0)})")
    
    # Skills & Gap Analysis
    doc.add_heading("3. Skills & Gap Analysis", level=1)
    doc.add_paragraph(f"Critical gaps: {summary.get('critical_gaps', 0)}")
    doc.add_paragraph(f"Internally fillable roles: {summary.get('fillable_internally', 0)}")
    
    # Talent Strategy
    doc.add_heading("4. Talent Strategy (BBBA)", level=1)
    bbba = summary.get("bbba_summary", {})
    doc.add_paragraph(f"Build (reskill internally): {bbba.get('build', 0)} roles")
    doc.add_paragraph(f"Buy (external hire): {bbba.get('buy', 0)} roles")
    doc.add_paragraph(f"Borrow (contract): {bbba.get('borrow', 0)} roles")
    doc.add_paragraph(f"Automate: {bbba.get('automate', 0)} roles")
    doc.add_paragraph(f"Total investment: ${bbba.get('total_investment', 0):,.0f}")
    
    # Headcount
    doc.add_heading("5. Headcount Planning", level=1)
    wf = summary.get("headcount_waterfall", {})
    doc.add_paragraph(f"Starting headcount: {wf.get('starting_headcount', 0)}")
    doc.add_paragraph(f"Target headcount: {wf.get('target_headcount', 0)}")
    doc.add_paragraph(f"Net change: {wf.get('net_change', 0)} ({wf.get('net_change_pct', 0)}%)")
    
    # Manager Capability
    doc.add_heading("6. Manager Capability", level=1)
    mgr = summary.get("manager_summary", {})
    doc.add_paragraph(f"Transformation Champions: {mgr.get('champions', 0)}")
    doc.add_paragraph(f"Needs Development: {mgr.get('needs_development', 0)}")
    doc.add_paragraph(f"Flight Risk: {mgr.get('flight_risk', 0)}")
    
    # Change Management
    doc.add_heading("7. Change Management", level=1)
    doc.add_paragraph(f"High-risk segment: {summary.get('high_risk_pct', 0)}% of workforce")
    
    # Reskilling
    doc.add_heading("8. Reskilling Program", level=1)
    reskill = summary.get("reskilling_summary", {})
    doc.add_paragraph(f"Employees requiring reskilling: {reskill.get('total_employees', 0)}")
    doc.add_paragraph(f"Average duration: {reskill.get('avg_months', 0)} months")
    doc.add_paragraph(f"Total reskilling investment: ${reskill.get('total_investment', 0):,.0f}")
    
    # Investment Summary
    doc.add_heading("9. Investment Summary", level=1)
    total = float(bbba.get('total_investment', 0)) + float(reskill.get('total_investment', 0))
    doc.add_paragraph(f"Talent sourcing: ${bbba.get('total_investment', 0):,.0f}")
    doc.add_paragraph(f"Reskilling: ${reskill.get('total_investment', 0):,.0f}")
    doc.add_paragraph(f"Total estimated investment: ${total:,.0f}")
    
    # Next Steps
    doc.add_heading("10. Recommended Next Steps", level=1)
    steps = [
        "Validate skill gap dispositions with function heads",
        "Begin reskilling programs for high-priority employees",
        "Engage flight-risk managers with executive coaching",
        "Launch change champion network in high-risk segments",
        "Finalize headcount plan with CHRO and CFO",
    ]
    for s in steps:
        doc.add_paragraph(s, style="List Bullet")
    
    # Save to buffer
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    
    from fastapi.responses import StreamingResponse
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=AI_Transformation_Report_{model_id}.docx"}
    )


@app.get("/api/export/template/{template_type}")
async def download_template(template_type: str):
    """Download Excel template with correct column headers."""
    import io
    from openpyxl import Workbook
    from fastapi.responses import StreamingResponse
    
    wb = Workbook()
    ws = wb.active
    
    templates = {
        "workforce": {
            "sheet_name": "Workforce",
            "headers": ["Employee ID", "Name", "Job Title", "Function", "Sub-Function", "Career Level", "Career Track", "Manager ID", "Compensation", "Tenure", "Location"],
            "samples": [
                ["EMP001", "Sarah Chen", "Data Analyst", "Technology", "Analytics", "L3", "IC", "EMP010", 85000, 3, "New York"],
                ["EMP002", "James Rodriguez", "Senior Data Analyst", "Technology", "Analytics", "L4", "IC", "EMP010", 105000, 5, "New York"],
                ["EMP003", "Maria Thompson", "Financial Analyst", "Finance", "FP&A", "L3", "IC", "EMP011", 90000, 4, "Chicago"],
                ["EMP004", "David Kim", "Operations Coordinator", "Operations", "Supply Chain", "L2", "IC", "EMP012", 65000, 2, "Chicago"],
                ["EMP010", "Rachel Anderson", "VP Technology", "Technology", "Leadership", "L5", "Manager", "EMP020", 185000, 8, "San Francisco"],
            ],
            "notes": "Required for: Workforce Snapshot, Job Architecture, Headcount Planning. Career Level: L1-L6. Career Track: IC or Manager.",
        },
        "work_design": {
            "sheet_name": "Work Design",
            "headers": ["Job Title", "Task Name", "Est Hours/Week", "Current Time Spent %", "Task Type", "Logic", "Interaction", "AI Impact", "Primary Skill", "Secondary Skill"],
            "samples": [
                ["Data Analyst", "Data extraction & cleaning", 10, 25, "Repetitive", "Deterministic", "Independent", "High", "Data Analysis", "Python/SQL"],
                ["Data Analyst", "Report generation & formatting", 8, 20, "Repetitive", "Deterministic", "Independent", "High", "Data Analysis", "Communication"],
                ["Data Analyst", "Ad-hoc analysis for stakeholders", 6, 15, "Variable", "Probabilistic", "Interactive", "Moderate", "Critical Thinking", "Data Analysis"],
                ["Data Analyst", "Stakeholder presentations", 4, 10, "Variable", "Judgment-heavy", "Collaborative", "Low", "Communication", "Stakeholder Mgmt"],
                ["Financial Analyst", "Monthly close reconciliation", 12, 30, "Repetitive", "Deterministic", "Independent", "High", "Data Analysis", "Process Automation"],
            ],
            "notes": "Required for: Work Design Lab, AI Opportunity Scan, Impact Simulator. Task Type: Repetitive/Variable. Logic: Deterministic/Probabilistic/Judgment-heavy. Interaction: Independent/Interactive/Collaborative. AI Impact: High/Moderate/Low. Time Spent % should sum to 100 per job.",
        },
        "skills": {
            "sheet_name": "Skills",
            "headers": ["Employee ID", "Name", "Skill", "Proficiency"],
            "samples": [
                ["EMP001", "Sarah Chen", "Data Analysis", 4],
                ["EMP001", "Sarah Chen", "Python/SQL", 3],
                ["EMP001", "Sarah Chen", "AI/ML Tools", 2],
                ["EMP001", "Sarah Chen", "Communication", 3],
                ["EMP002", "James Rodriguez", "Data Analysis", 4],
                ["EMP002", "James Rodriguez", "Leadership", 3],
                ["EMP003", "Maria Thompson", "Data Analysis", 3],
                ["EMP003", "Maria Thompson", "Critical Thinking", 3],
            ],
            "notes": "Required for: Skills Inventory, Gap Analysis, Adjacency Map. Proficiency scale: 1=Novice, 2=Developing, 3=Proficient, 4=Expert. One row per employee per skill.",
        },
        "org_structure": {
            "sheet_name": "Org Structure",
            "headers": ["Employee ID", "Name", "Job Title", "Function", "Career Level", "Manager ID", "Span of Control", "Cost Center", "Location"],
            "samples": [
                ["EMP001", "Sarah Chen", "Data Analyst", "Technology", "L3", "EMP010", 0, "CC100", "New York"],
                ["EMP010", "Rachel Anderson", "VP Technology", "Technology", "L5", "EMP020", 5, "CC100", "San Francisco"],
                ["EMP020", "Andrew Scott", "CEO", "Executive", "L6", "", 5, "CC000", "New York"],
            ],
            "notes": "Required for: Org Design Studio, Manager Capability. Span of Control = number of direct reports. Manager ID links to another employee's Employee ID.",
        },
        "ai_readiness": {
            "sheet_name": "AI Readiness",
            "headers": ["Employee ID", "Name", "Function", "AI Awareness", "Tool Adoption", "Data Literacy", "Change Openness", "AI Collaboration"],
            "samples": [
                ["EMP001", "Sarah Chen", "Technology", 4, 3, 4, 3, 3],
                ["EMP002", "James Rodriguez", "Technology", 4, 4, 4, 3, 4],
                ["EMP003", "Maria Thompson", "Finance", 3, 2, 3, 3, 2],
                ["EMP004", "David Kim", "Operations", 2, 1, 2, 2, 1],
                ["EMP010", "Rachel Anderson", "Technology", 4, 3, 3, 4, 3],
            ],
            "notes": "Required for: AI Readiness Assessment, Change Readiness. Score each dimension 1-5. 1=No awareness, 2=Basic, 3=Intermediate, 4=Advanced, 5=Expert.",
        },
        "manager_capability": {
            "sheet_name": "Manager Capability",
            "headers": ["Employee ID", "Name", "Function", "Direct Reports", "Leading Through Ambiguity", "AI Tool Fluency", "Coaching & Reskilling", "Communication of Change"],
            "samples": [
                ["EMP010", "Rachel Anderson", "Technology", 5, 4, 4, 3, 4],
                ["EMP011", "Thomas Harris", "Finance", 3, 3, 2, 3, 3],
                ["EMP012", "Catherine Clark", "Operations", 4, 2, 1, 2, 2],
                ["EMP013", "Steven Wright", "HR", 3, 3, 3, 4, 4],
            ],
            "notes": "Required for: Manager Capability Assessment, Manager Development. Only include people managers (those with direct reports). Score each dimension 1-5.",
        },
        "operating_model": {
            "sheet_name": "Operating Model",
            "headers": ["Function", "Capability", "Current Maturity", "Target Maturity", "Service Model", "Owner"],
            "samples": [
                ["Finance", "Accounting & Close", 3, 4, "Shared", "VP Finance"],
                ["Finance", "FP&A", 2, 4, "Embedded", "VP Finance"],
                ["Finance", "Treasury", 2, 3, "Shared", "VP Finance"],
                ["Technology", "Data & Analytics", 3, 5, "Platform", "VP Technology"],
                ["Technology", "Infrastructure", 3, 4, "Shared", "VP Technology"],
            ],
            "notes": "Required for: Operating Model Lab. Maturity: 1=Ad Hoc, 2=Repeatable, 3=Defined, 4=Managed, 5=Optimized. Service Model: Shared/Embedded/CoE/Outsourced/Platform.",
        },
        "change_management": {
            "sheet_name": "Change Management",
            "headers": ["Employee ID", "Name", "Function", "Change Readiness", "Impact Level", "Support Needed", "Preferred Communication", "Manager ID"],
            "samples": [
                ["EMP001", "Sarah Chen", "Technology", 4, "High", "Light", "Digital/Email", "EMP010"],
                ["EMP003", "Maria Thompson", "Finance", 3, "Moderate", "Moderate", "Town Hall", "EMP011"],
                ["EMP004", "David Kim", "Operations", 2, "High", "Intensive", "1:1 Meeting", "EMP012"],
                ["EMP009", "Jennifer Garcia", "Legal", 3, "Low", "Light", "Newsletter", "EMP014"],
            ],
            "notes": "Required for: Change Readiness, Change Planner. Change Readiness: 1-5. Impact Level: High/Moderate/Low. Support Needed: Intensive/Moderate/Light.",
        },
    }
    
    tmpl = templates.get(template_type)
    if not tmpl:
        return {"error": f"Unknown template: {template_type}. Available: {list(templates.keys())}"}
    
    ws.title = tmpl["sheet_name"]
    ws.append(tmpl["headers"])
    
    # Add sample rows
    samples = tmpl.get("samples", [tmpl.get("sample", [])])
    for sample in samples:
        ws.append(sample)
    
    # Style headers
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    header_fill = PatternFill(start_color="1A2340", end_color="1A2340", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    thin_border = Border(
        bottom=Side(style='thin', color='3B82F6'),
    )
    for cell in ws[1]:
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = thin_border
    
    # Style sample data rows
    sample_font = Font(color="888888", italic=True, size=10)
    for row in ws.iter_rows(min_row=2, max_row=len(samples)+1):
        for cell in row:
            cell.font = sample_font
    
    # Auto-width
    for col in ws.columns:
        max_len = max(len(str(cell.value or "")) for cell in col)
        ws.column_dimensions[col[0].column_letter].width = max(max_len + 4, 14)
    
    # Add notes sheet if available
    if tmpl.get("notes"):
        notes_ws = wb.create_sheet("Instructions")
        notes_ws.append(["Instructions for " + tmpl["sheet_name"]])
        notes_ws.append([])
        notes_ws.append([tmpl["notes"]])
        notes_ws.append([])
        notes_ws.append(["Column Definitions:"])
        for i, h in enumerate(tmpl["headers"]):
            notes_ws.append([f"  {h}", "— Required" if i < 3 else "— Recommended"])
        notes_ws.append([])
        notes_ws.append(["Tips:"])
        notes_ws.append(["  • Delete the sample rows before uploading your data"])
        notes_ws.append(["  • Do not change column header names"])
        notes_ws.append(["  • Keep the sheet name as-is"])
        notes_ws.column_dimensions['A'].width = 50
        notes_ws.column_dimensions['B'].width = 20
        notes_ws['A1'].font = Font(bold=True, size=12, color="3B82F6")
    
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={template_type}_template.xlsx"}
    )


# ═══════════════════════════════════════════════════
# PER-MODULE DATA EXPORTS
# ═══════════════════════════════════════════════════

@app.get("/api/export/module/{model_id}/{module_name}")
async def export_module_data(model_id: str, module_name: str):
    """Export module-specific data as Excel."""
    import io
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    from fastapi.responses import StreamingResponse
    
    wb = Workbook()
    ws = wb.active
    
    header_fill = PatternFill(start_color="1A2340", end_color="1A2340", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    
    def style_headers():
        for cell in ws[1]:
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center')
        for col in ws.columns:
            max_len = max(len(str(cell.value or "")) for cell in col)
            ws.column_dimensions[col[0].column_letter].width = max(max_len + 4, 14)
    
    if module_name == "snapshot":
        ws.title = "Workforce Overview"
        ds = store.datasets.get(model_id)
        if ds:
            wf = None
            for key in ds:
                if "workforce" in key.lower() or "employee" in key.lower():
                    wf = ds[key]; break
            if wf is not None and not wf.empty:
                ws.append(list(wf.columns))
                for _, row in wf.iterrows():
                    ws.append([str(v) for v in row.values])
                style_headers()
        
    elif module_name == "skills_inventory":
        inv = await get_skills_inventory(model_id)
        ws.title = "Skills Inventory"
        ws.append(["Employee", "Skill", "Proficiency", "Cluster"])
        records = inv.get("records", [])
        clusters = inv.get("clusters", {})
        skill_cluster = {}
        for cluster, skills in clusters.items():
            for s in skills:
                skill_cluster[s] = cluster
        for r in records:
            ws.append([r["employee"], r["skill"], r["proficiency"], skill_cluster.get(r["skill"], "Other")])
        style_headers()
    
    elif module_name == "skills_gap":
        gap = await get_skills_gap(model_id)
        ws.title = "Skills Gap Analysis"
        ws.append(["Skill", "Current Avg", "Target", "Target Source", "Delta", "Severity", "Employees Assessed"])
        for g in gap.get("gaps", []):
            ws.append([g["skill"], g["current_avg"], g["target"], g.get("target_source", ""), g["delta"], g["severity"], g["employees_assessed"]])
        style_headers()
    
    elif module_name == "adjacency":
        adj = await get_skills_adjacency(model_id)
        ws.title = "Adjacency Map"
        ws.append(["Target Role", "Candidate", "Adjacency %", "Matching Skills", "Gap Skills", "Reskill Months"])
        for a in adj.get("adjacencies", []):
            for c in a.get("top_candidates", []):
                ws.append([a["target_role"], c["employee"], c["adjacency_pct"], ", ".join(c.get("matching_skills", [])), ", ".join(c.get("gap_skills", [])), c.get("reskill_months", 0)])
        style_headers()
    
    elif module_name == "bbba":
        bbba = await get_bbba(model_id)
        ws.title = "BBBA Decisions"
        ws.append(["Role", "Disposition", "Reason", "Strong Candidates", "Reskillable", "FTE Needed", "Cost per FTE", "Total Cost", "Timeline (months)"])
        for r in bbba.get("roles", []):
            ws.append([r["role"], r["disposition"], r["reason"], r["strong_candidates"], r["reskillable_candidates"], r["fte_needed"], r["cost_per_fte"], r["total_cost"], r["timeline_months"]])
        style_headers()
    
    elif module_name == "headcount":
        hc = await get_headcount_plan(model_id)
        ws.title = "Headcount Plan"
        wf = hc.get("waterfall", {})
        ws.append(["Metric", "Value"])
        for k, v in wf.items():
            ws.append([k.replace("_", " ").title(), v])
        ws.append([])
        ws.append(["Department", "Current FTE", "Eliminated", "Redeployed", "New Hires", "Future FTE", "% Change"])
        for d in hc.get("departments", []):
            ws.append([d["department"], d["current_fte"], d["eliminated"], d["redeployed"], d["new_hires"], d["future_fte"], d["pct_change"]])
        style_headers()
    
    elif module_name == "readiness":
        rd = await get_readiness_assessment(model_id)
        ws.title = "AI Readiness"
        dims = rd.get("dimensions", [])
        ws.append(["Employee", *dims, "Average", "Band"])
        for ind in rd.get("individuals", []):
            ws.append([ind["employee"], *[ind["scores"].get(d, 0) for d in dims], ind["average"], ind["band"]])
        style_headers()
    
    elif module_name == "manager_capability":
        mgr = await get_manager_capability(model_id)
        ws.title = "Manager Capability"
        dims = mgr.get("dimensions", [])
        ws.append(["Manager", "Direct Reports", *dims, "Average", "Category", "Team Readiness"])
        for m in mgr.get("managers", []):
            ws.append([m["manager"], m["direct_reports"], *[m["scores"].get(d, 0) for d in dims], m["average"], m["category"], m["team_readiness_avg"]])
        style_headers()
    
    elif module_name == "reskilling":
        rsk = await get_reskilling_pathways(model_id)
        ws.title = "Reskilling Pathways"
        ws.append(["Employee", "Target Role", "Priority", "Readiness Score", "Readiness Band", "Total Months", "Estimated Cost", "Skills to Develop"])
        for p in rsk.get("pathways", []):
            skills = ", ".join(s["skill"] for s in p.get("skills_to_develop", []))
            ws.append([p["employee"], p["target_role"], p["priority"], p["readiness_score"], p["readiness_band"], p["total_months"], p["estimated_cost"], skills])
        style_headers()
    
    elif module_name == "marketplace":
        mp = await get_talent_marketplace(model_id)
        ws.title = "Talent Marketplace"
        ws.append(["Target Role", "Fill Recommendation", "Candidate", "Composite Score", "Adjacency %", "Readiness", "Reskill Months", "Has Pathway", "Pathway Cost"])
        for m in mp.get("marketplace", []):
            for c in m.get("candidates", []):
                ws.append([m["target_role"], m["fill_recommendation"], c["employee"], c["composite_score"], c["adjacency_pct"], c["readiness_score"], c["reskill_months"], c["has_pathway"], c["pathway_cost"]])
        style_headers()
    
    elif module_name == "change_readiness":
        cr = await get_change_readiness(model_id)
        ws.title = "Change Readiness"
        ws.append(["Employee", "Readiness Score", "Impact Score", "Band", "Quadrant", "Intervention"])
        for qkey, qdata in cr.get("quadrants", {}).items():
            for emp in qdata.get("employees", []):
                ws.append([emp["employee"], emp["readiness"], emp["impact"], emp["band"], qdata["label"], qdata["action"]])
        style_headers()
    
    elif module_name == "manager_development":
        md = await get_manager_development(model_id)
        ws.title = "Manager Development"
        ws.append(["Manager", "Category", "Score", "Direct Reports", "Role in Change", "Weak Dimensions", "Total Weeks", "Total Cost"])
        for t in md.get("tracks", []):
            ws.append([t["manager"], t["category"], t["average_score"], t["direct_reports"], t["role_in_change"], ", ".join(t["weak_dimensions"]), t["total_weeks"], t["total_cost"]])
        style_headers()
    
    else:
        ws.append(["Error", f"Unknown module: {module_name}"])
    
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={module_name}_{model_id}.xlsx"}
    )
